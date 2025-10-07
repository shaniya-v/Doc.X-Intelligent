import os
import json
import base64
import io
import uuid
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional
import logging

import pandas as pd
import numpy as np
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import requests

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    from PIL import Image
    import pytesseract
    from openpyxl import load_workbook
    import docx
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Document processing libraries not available: {e}")
    DOCUMENT_PROCESSING_AVAILABLE = False

# Database
from supabase import create_client, Client

# Import enhanced classifiers
from rag_classifier import RAGDepartmentClassifier
from multi_department_classifier import MultiDepartmentClassifier

# Document utilities
def calculate_content_hash(content: str) -> str:
    """Calculate SHA-256 hash of document content for duplicate detection"""
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

def check_document_exists(filename: str, content: str, user_id: str = None) -> Dict[str, Any]:
    """Check if document already exists by filename and content hash"""
    try:
        content_hash = calculate_content_hash(content)
        
        # Check by filename
        filename_result = supabase.table('documents').select('*').eq('title', filename).execute()
        
        # Check by content hash in metadata (since content_hash column may not exist)
        all_docs_result = supabase.table('documents').select('*').execute()
        
        existing_by_name = filename_result.data if filename_result.data else []
        existing_by_content = []
        
        # Check content hash in metadata
        for doc in (all_docs_result.data or []):
            doc_metadata = doc.get('metadata', {})
            if doc_metadata.get('content_hash') == content_hash:
                existing_by_content.append(doc)
        
        return {
            'exists_by_name': len(existing_by_name) > 0,
            'exists_by_content': len(existing_by_content) > 0,
            'existing_name_docs': existing_by_name,
            'existing_content_docs': existing_by_content,
            'content_hash': content_hash,
            'is_duplicate': len(existing_by_content) > 0  # True duplicate if content matches
        }
        
    except Exception as e:
        logger.error(f"Error checking document existence: {e}")
        return {
            'exists_by_name': False,
            'exists_by_content': False,
            'existing_name_docs': [],
            'existing_content_docs': [],
            'content_hash': calculate_content_hash(content),
            'is_duplicate': False
        }

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files"""
    try:
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return f"PDF content from {file_path} - PDF processing not available"
        
        text = ""
        try:
            # Try with pdfplumber first (better for complex layouts)
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
                return f"PDF document - unable to extract text content"
        
        return text.strip() if text.strip() else f"PDF document - no readable text content found"
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"PDF document - extraction error: {str(e)}"

def extract_text_from_word(file_path: str) -> str:
    """Extract text from Word documents"""
    try:
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return f"Word document content from {file_path} - Word processing not available"
        
        try:
            import docx
            doc = docx.Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            content = "\n".join(text)
            return content.strip() if content.strip() else f"Word document - no readable content found"
            
        except Exception as e:
            logger.warning(f"Word extraction failed: {e}")
            return f"Word document - unable to extract text content"
        
    except Exception as e:
        logger.error(f"Error extracting text from Word: {e}")
        return f"Word document - extraction error: {str(e)}"

def extract_text_from_excel(file_path: str) -> str:
    """Extract text from Excel files"""
    try:
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return f"Excel content from {file_path} - Excel processing not available"
        
        try:
            import pandas as pd
            
            # Try to read Excel file
            excel_data = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            
            text = []
            for sheet_name, df in excel_data.items():
                text.append(f"Sheet: {sheet_name}")
                text.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                text.append(f"Columns: {', '.join(df.columns.astype(str))}")
                
                # Add sample data
                if len(df) > 0:
                    text.append("Sample data:")
                    text.append(df.head(3).to_string(index=False))
                text.append("")
            
            content = "\n".join(text)
            return content.strip() if content.strip() else f"Excel document - no readable content found"
            
        except Exception as e:
            logger.warning(f"Excel extraction failed: {e}")
            return f"Excel document - unable to extract content"
        
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {e}")
        return f"Excel document - extraction error: {str(e)}"

def extract_text_from_csv(file_path: str) -> str:
    """Extract text from CSV files"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        
        text = []
        text.append(f"CSV Data: {len(df)} rows, {len(df.columns)} columns")
        text.append(f"Columns: {', '.join(df.columns.astype(str))}")
        text.append("Sample data:")
        text.append(df.head(5).to_string(index=False))
        
        return "\n".join(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {e}")
        return f"CSV document - extraction error: {str(e)}"

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app, origins=[
    "http://localhost:3000", 
    "http://localhost:3001", 
    "http://127.0.0.1:3000", 
    "http://127.0.0.1:3001"
])

# Initialize Supabase client
supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_ANON_KEY")
supabase: Client = create_client(supabase_url, supabase_key)

# OpenRouter configuration
openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

class DocumentProcessor:
    """Advanced document processor with OCR, PDF, Excel, and image processing"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'txt', 'json', 'xlsx', 'xls', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'tiff', 'bmp']
        # Configure Tesseract OCR path if needed (Windows)
        # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    def process_binary_file(self, binary_data: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Process binary file data and extract text content"""
        try:
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            
            print(f"🔍 Processing {filename} ({mime_type}, {len(binary_data)} bytes)")
            
            extracted_text = ""
            metadata = {
                "filename": filename,
                "mime_type": mime_type,
                "file_size": len(binary_data),
                "file_extension": file_ext,
                "processing_method": "unknown"
            }
            
            # PDF Processing
            if file_ext == 'pdf' or 'pdf' in mime_type.lower():
                extracted_text, metadata = self.extract_pdf_content(binary_data, metadata)
            
            # Excel Processing
            elif file_ext in ['xlsx', 'xls'] or 'excel' in mime_type.lower() or 'spreadsheet' in mime_type.lower():
                extracted_text, metadata = self.extract_excel_content(binary_data, metadata)
            
            # Word Document Processing
            elif file_ext in ['docx', 'doc'] or 'word' in mime_type.lower():
                extracted_text, metadata = self.extract_word_content(binary_data, metadata)
            
            # Image Processing with OCR
            elif file_ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp'] or 'image' in mime_type.lower():
                extracted_text, metadata = self.extract_image_content(binary_data, metadata)
            
            # Plain text
            elif file_ext == 'txt' or 'text' in mime_type.lower():
                extracted_text = binary_data.decode('utf-8', errors='ignore')
                metadata["processing_method"] = "text_decode"
            
            else:
                # Try to decode as text as fallback
                try:
                    extracted_text = binary_data.decode('utf-8', errors='ignore')
                    metadata["processing_method"] = "fallback_text_decode"
                except:
                    extracted_text = f"Binary file: {filename} (unsupported format)"
                    metadata["processing_method"] = "unsupported"
            
            # Analyze extracted text
            if extracted_text.strip():
                language = self.detect_language(extracted_text)
                word_count = len(extracted_text.split())
                
                return {
                    "text": extracted_text,
                    "language": language,
                    "word_count": word_count,
                    "metadata": metadata,
                    "success": True
                }
            else:
                return {
                    "text": f"File processed but no text content extracted from {filename}",
                    "language": "unknown",
                    "word_count": 0,
                    "metadata": metadata,
                    "success": False
                }
                
        except Exception as e:
            logger.error(f"Error processing binary file {filename}: {str(e)}")
            return {
                "text": f"Error processing {filename}: {str(e)}",
                "language": "unknown", 
                "word_count": 0,
                "metadata": {"error": str(e), "filename": filename},
                "success": False
            }
    
    def extract_pdf_content(self, binary_data: bytes, metadata: dict) -> tuple:
        """Extract text from PDF with multiple methods including tables"""
        extracted_text = ""
        
        try:
            # Method 1: Try pdfplumber (better for complex layouts and tables)
            if DOCUMENT_PROCESSING_AVAILABLE:
                try:
                    pdf_file = io.BytesIO(binary_data)
                    with pdfplumber.open(pdf_file) as pdf:
                        content_parts = []
                        tables_found = 0
                        
                        for page_num, page in enumerate(pdf.pages):
                            page_content = [f"Page {page_num + 1}:"]
                            
                            # Extract regular text
                            page_text = page.extract_text()
                            if page_text:
                                page_content.append(page_text)
                            
                            # Extract tables
                            tables = page.extract_tables()
                            for table_num, table in enumerate(tables):
                                if table:
                                    page_content.append(f"\nTable {table_num + 1} (Page {page_num + 1}):")
                                    for row in table:
                                        if row:
                                            row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                            if row_text.strip():
                                                page_content.append(row_text)
                                    tables_found += 1
                            
                            if len(page_content) > 1:  # More than just page header
                                content_parts.append("\n".join(page_content))
                        
                        if content_parts:
                            extracted_text = "\n\n".join(content_parts)
                            metadata["processing_method"] = "pdfplumber_enhanced"
                            metadata["pages_processed"] = len(pdf.pages)
                            metadata["tables_extracted"] = tables_found
                            return extracted_text, metadata
                except Exception as e:
                    print(f"pdfplumber failed: {e}")
            
            # Method 2: Try PyPDF2 as fallback
            if DOCUMENT_PROCESSING_AVAILABLE and not extracted_text:
                try:
                    pdf_file = io.BytesIO(binary_data)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text_parts = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"Page {page_num + 1}:\n{page_text}")
                    
                    if text_parts:
                        extracted_text = "\n\n".join(text_parts)
                        metadata["processing_method"] = "pypdf2"
                        metadata["pages_processed"] = len(pdf_reader.pages)
                        return extracted_text, metadata
                except Exception as e:
                    print(f"PyPDF2 failed: {e}")
            
            # Method 3: If text extraction fails, indicate it might be image-based
            if not extracted_text:
                extracted_text = "PDF processed but limited text found. This may be an image-based PDF that requires OCR processing."
                metadata["processing_method"] = "pdf_limited_text"
            
        except Exception as e:
            print(f"PDF processing error: {e}")
            metadata["processing_error"] = str(e)
        
        if not extracted_text:
            extracted_text = "PDF file processed but text extraction failed"
            metadata["processing_method"] = "pdf_failed"
        
        return extracted_text, metadata
    
    def extract_excel_content(self, binary_data: bytes, metadata: dict) -> tuple:
        """Extract content from Excel files"""
        extracted_text = ""
        
        try:
            if DOCUMENT_PROCESSING_AVAILABLE:
                excel_file = io.BytesIO(binary_data)
                
                # Try openpyxl first (for .xlsx)
                try:
                    workbook = load_workbook(excel_file, data_only=True)
                    sheets_content = []
                    
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        sheet_content = [f"Sheet: {sheet_name}"]
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                            if row_text.strip():
                                sheet_content.append(row_text)
                        
                        sheets_content.append("\n".join(sheet_content))
                    
                    extracted_text = "\n\n".join(sheets_content)
                    metadata["processing_method"] = "openpyxl"
                    metadata["sheets_processed"] = len(workbook.sheetnames)
                    
                except Exception as e:
                    print(f"openpyxl failed: {e}")
                    
                    # Try pandas as fallback
                    try:
                        excel_file.seek(0)
                        excel_data = pd.read_excel(excel_file, sheet_name=None)
                        sheets_content = []
                        
                        for sheet_name, df in excel_data.items():
                            sheet_content = f"Sheet: {sheet_name}\n"
                            sheet_content += df.to_string(index=False)
                            sheets_content.append(sheet_content)
                        
                        extracted_text = "\n\n".join(sheets_content)
                        metadata["processing_method"] = "pandas"
                        metadata["sheets_processed"] = len(excel_data)
                        
                    except Exception as e2:
                        print(f"pandas Excel processing failed: {e2}")
                        extracted_text = f"Excel file detected but processing failed: {str(e2)}"
                        metadata["processing_method"] = "excel_failed"
            
        except Exception as e:
            print(f"Excel processing error: {e}")
            metadata["processing_error"] = str(e)
            extracted_text = f"Excel processing error: {str(e)}"
        
        return extracted_text, metadata
    
    def extract_word_content(self, binary_data: bytes, metadata: dict) -> tuple:
        """Extract content from Word documents including tables"""
        extracted_text = ""
        
        try:
            if DOCUMENT_PROCESSING_AVAILABLE:
                doc_file = io.BytesIO(binary_data)
                doc = docx.Document(doc_file)
                
                content_parts = []
                
                # Extract paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                if paragraphs:
                    content_parts.append("Document Content:")
                    content_parts.extend(paragraphs)
                
                # Extract tables
                tables_content = []
                for i, table in enumerate(doc.tables):
                    table_content = [f"\nTable {i+1}:"]
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                        if row_data:
                            table_content.append(" | ".join(row_data))
                    
                    if len(table_content) > 1:  # More than just the header
                        tables_content.extend(table_content)
                
                if tables_content:
                    content_parts.extend(tables_content)
                
                extracted_text = "\n".join(content_parts)
                metadata["processing_method"] = "python-docx-enhanced"
                metadata["paragraphs_processed"] = len(paragraphs)
                metadata["tables_processed"] = len(doc.tables)
                
        except Exception as e:
            print(f"Word processing error: {e}")
            metadata["processing_error"] = str(e)
            extracted_text = f"Word document processing error: {str(e)}"
            metadata["processing_method"] = "word_failed"
        
        return extracted_text, metadata
    
    def extract_image_content(self, binary_data: bytes, metadata: dict) -> tuple:
        """Extract text from images using OCR"""
        extracted_text = ""
        
        try:
            if DOCUMENT_PROCESSING_AVAILABLE:
                image = Image.open(io.BytesIO(binary_data))
                
                # Perform OCR
                extracted_text = pytesseract.image_to_string(image, lang='eng+mal')  # English + Malayalam
                
                if extracted_text.strip():
                    metadata["processing_method"] = "tesseract_ocr"
                    metadata["image_size"] = image.size
                    metadata["image_mode"] = image.mode
                else:
                    extracted_text = "Image processed but no text detected via OCR"
                    metadata["processing_method"] = "ocr_no_text"
                
        except Exception as e:
            print(f"Image OCR error: {e}")
            metadata["processing_error"] = str(e)
            extracted_text = f"Image OCR processing error: {str(e)}"
            metadata["processing_method"] = "ocr_failed"
        
        return extracted_text, metadata
    
    def extract_text_from_content(self, content: str, content_type: str = "text") -> Dict[str, Any]:
        """Extract and analyze text content (legacy method for backward compatibility)"""
        try:
            result = {
                "text": content,
                "language": self.detect_language(content),
                "word_count": len(content.split()),
                "metadata": {"content_type": content_type}
            }
            return result
        except Exception as e:
            logger.error(f"Error processing content: {str(e)}")
            return {"error": str(e)}
    
    def detect_language(self, text: str) -> str:
        """Detect language of text (English/Malayalam/Mixed)"""
        if not text or len(text.strip()) < 10:
            return "unknown"
        
        try:
            # Check for Malayalam Unicode range
            malayalam_chars = sum(1 for char in text if '\u0d00' <= char <= '\u0d7f')
            english_chars = sum(1 for char in text if char.isascii() and char.isalpha())
            total_chars = malayalam_chars + english_chars
            
            if total_chars == 0:
                return "unknown"
            
            malayalam_ratio = malayalam_chars / total_chars
            english_ratio = english_chars / total_chars
            
            if malayalam_ratio > 0.3 and english_ratio > 0.3:
                return "mixed_en_ml"
            elif malayalam_ratio > 0.5:
                return "malayalam"
            elif english_ratio > 0.7:
                return "english"
            else:
                return "mixed"
        except Exception as e:
            logger.error(f"Language detection error: {str(e)}")
            return "unknown"

class SmartRouter:
    """Intelligent document routing based on content analysis"""
    
    def __init__(self):
        self.department_keywords = {
            "engineering": [
                "maintenance", "repair", "technical", "equipment", "machinery", 
                "infrastructure", "construction", "mechanical", "electrical", "metro",
                "train", "track", "signal", "platform", "station", "rail",
                # Malayalam keywords
                "അറ്റകുറ്റപ്പണി", "സാങ്കേതിക", "യന്ത്രസാമഗ്രി", "മെട്രോ", "ട്രെയിൻ"
            ],
            "finance": [
                "budget", "payment", "invoice", "cost", "expense", "financial",
                "accounting", "procurement", "purchase", "money", "tender", "contract",
                # Malayalam keywords  
                "ബജറ്റ്", "പണം", "വിലക", "ചെലവ്", "കരാർ"
            ],
            "hr": [
                "employee", "staff", "personnel", "leave", "attendance", "training",
                "recruitment", "salary", "benefits", "policy", "transfer",
                # Malayalam keywords
                "ജീവനക്കാർ", "ശമ്പളം", "അവധി", "പരിശീലനം", "നയം"
            ],
            "admin": [
                "administrative", "office", "supplies", "facility", "security",
                "general", "coordination", "management", "policy", "meeting",
                # Malayalam keywords
                "ഭരണ", "ഓഫീസ്", "സുരക്ഷ", "മീറ്റിംഗ്"
            ],
            "safety": [
                "safety", "hazard", "risk", "accident", "emergency", "incident",
                "compliance", "audit", "inspection", "protocol", "fire", "evacuation",
                # Malayalam keywords
                "സുരക്ഷ", "അപകടം", "അപകടസാധ്യത", "അഗ്നി"
            ],
            "operations": [
                "schedule", "timetable", "service", "passenger", "operation", "control",
                "monitoring", "performance", "delay", "disruption",
                # Malayalam keywords
                "സമയക്രമം", "സർവീസ്", "യാത്രക്കാർ", "പ്രവർത്തനം"
            ]
        }
    
    def route_document(self, content: str, metadata: Dict) -> Dict[str, Any]:
        """Route document to appropriate department based on content analysis"""
        
        # Convert content to lowercase for matching
        content_lower = content.lower()
        
        # Calculate department scores
        department_scores = {}
        total_words = len(content.split())
        
        for dept, keywords in self.department_keywords.items():
            score = 0
            matched_keywords = []
            
            for keyword in keywords:
                if keyword.lower() in content_lower:
                    # Weight longer keywords more
                    weight = len(keyword.split())
                    score += weight
                    matched_keywords.append(keyword)
            
            if score > 0:
                department_scores[dept] = {
                    "score": score,
                    "keywords": matched_keywords,
                    "relevance": min((score / max(total_words, 1)) * 100, 100)
                }
        
        # Determine priority based on keywords
        priority = self.determine_priority(content_lower)
        
        # Get the department with highest score
        if department_scores:
            best_dept = max(department_scores, key=lambda x: department_scores[x]["score"])
            confidence = department_scores[best_dept]["relevance"]
            reasoning = f"Matched keywords: {', '.join(department_scores[best_dept]['keywords'][:3])}"
        else:
            best_dept = "admin"  # Default department
            confidence = 10
            reasoning = "No specific keywords found, assigned to admin for review"
        
        return {
            "assigned_department": best_dept,
            "confidence": min(confidence, 100),
            "priority": priority,
            "department_scores": {k: v["score"] for k, v in department_scores.items()},
            "reasoning": reasoning,
            "matched_keywords": department_scores.get(best_dept, {}).get("keywords", [])
        }
    
    def determine_priority(self, content: str) -> str:
        """Determine document priority based on content"""
        urgent_keywords = [
            "urgent", "emergency", "critical", "immediate", "asap", "breakdown",
            "failure", "accident", "incident", "disruption", "fire", "evacuation",
            # Malayalam urgent keywords
            "തുരന്തം", "അടിയന്തര", "ഗുരുതര", "അപകടം"
        ]
        
        high_keywords = [
            "important", "priority", "deadline", "schedule", "maintenance",
            "repair", "inspection", "safety", "compliance",
            # Malayalam important keywords
            "പ്രധാന", "സമയം", "സുരക്ഷ"
        ]
        
        if any(keyword in content for keyword in urgent_keywords):
            return "urgent"
        elif any(keyword in content for keyword in high_keywords):
            return "high"
        else:
            return "normal"

class AIAnalyzer:
    """AI-powered document analysis using OpenRouter"""
    
    def __init__(self):
        self.headers = {
            "Authorization": f"Bearer {openrouter_api_key}",
            "Content-Type": "application/json"
        }
    
    def analyze_document(self, content: str, context: Dict = None) -> Dict[str, Any]:
        """Analyze document content using AI"""
        try:
            prompt = f"""
            Analyze the following KMRL (Kochi Metro Rail Limited) document and provide:
            1. Main topic/subject
            2. Department it should be routed to (engineering, finance, hr, admin, safety, operations)
            3. Priority level (urgent, high, normal, low)
            4. Key action items or requirements
            5. Language detected (english, malayalam, mixed)
            
            Document content:
            {content[:1000]}  # Limit content for API call
            
            Provide response in JSON format.
            """
            
            payload = {
                "model": "meta-llama/llama-3.2-3b-instruct:free",
                "messages": [
                    {"role": "system", "content": "You are an expert document analyzer for KMRL operations."},
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": 500,
                "temperature": 0.3
            }
            
            response = requests.post(OPENROUTER_URL, headers=self.headers, json=payload)
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
                
                # Try to parse JSON response
                try:
                    analysis = json.loads(ai_response)
                    return analysis
                except json.JSONDecodeError:
                    # If not JSON, return structured response
                    return {
                        "analysis": ai_response,
                        "confidence": 70,
                        "ai_powered": True
                    }
            else:
                logger.error(f"OpenRouter API error: {response.status_code}")
                return {"error": "AI analysis failed", "fallback": True}
                
        except Exception as e:
            logger.error(f"AI analysis error: {str(e)}")
            return {"error": str(e), "fallback": True}

# Initialize processors
doc_processor = DocumentProcessor()
smart_router = SmartRouter()
ai_analyzer = AIAnalyzer()

# Initialize RAG classifier for advanced document routing
try:
    from rag_classifier import RAGDepartmentClassifier
    rag_classifier = RAGDepartmentClassifier()
    print("🤖 RAG Classifier initialized for webhook processing")
except Exception as e:
    print(f"⚠️ RAG Classifier initialization failed: {e}")
    rag_classifier = None

# Initialize multi-department classifier for complex documents
try:
    from multi_department_classifier import MultiDepartmentClassifier
    multi_dept_classifier = MultiDepartmentClassifier()
    print("🎯 Multi-Department Classifier initialized")
except Exception as e:
    print(f"⚠️ Multi-Department Classifier initialization failed: {e}")
    multi_dept_classifier = None

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy", 
        "service": "DOC.X Intelligent Backend",
        "features": [
            "Multi-format document processing",
            "Bilingual support (English/Malayalam)",
            "Smart department routing",
            "AI-powered analysis",
            "Supabase integration"
        ],
        "timestamp": datetime.now().isoformat()
    })

@app.route('/webhook/document', methods=['POST'])
def process_document_webhook():
    """Enhanced webhook for processing documents with binary file support - N8N ONLY"""
    try:
        # Check for N8N authentication
        auth_header = request.headers.get('Authorization')
        n8n_source = request.headers.get('X-N8N-Source')
        
        # Allow N8N or local testing
        is_n8n_request = (
            auth_header and 'n8n' in auth_header.lower() or
            n8n_source or
            request.headers.get('User-Agent', '').startswith('n8n') or
            request.remote_addr in ['127.0.0.1', 'localhost'] or
            request.args.get('source') == 'n8n'
        )
        
        if not is_n8n_request:
            return jsonify({
                "error": "Unauthorized - Documents can only be submitted via N8N workflow",
                "message": "Direct document submission is not allowed. Please use the N8N integration."
            }), 403
        
        # Handle both JSON and multipart form data
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle file uploads
            data = request.form.to_dict()
            files = request.files
        else:
            # Handle JSON data from N8N
            data = request.get_json() or {}
            files = {}
        
        logger.info(f"🔍 Received document processing request: {list(data.keys())}")
        
        # Debug: Print the actual data received
        print(f"🔍 Raw data content: {data.get('content', 'NO CONTENT FIELD')}")
        print(f"🔍 Raw data subject: {data.get('subject', 'NO SUBJECT FIELD')}")
        print(f"🔍 Available keys: {list(data.keys())}")
        
        # Extract document information with multiple possible field names
        document_data = {
            "source": data.get("source", "email"),
            "content": data.get("content", data.get("body", data.get("text", ""))),
            "subject": data.get("subject", data.get("title", data.get("filename", ""))),
            "sender": data.get("from", data.get("sender", "")),
            "timestamp": data.get("timestamp", datetime.now().isoformat()),
            "message_id": data.get("messageId", data.get("message_id", data.get("documentId", ""))),
            "attachments": data.get("attachments", []),
            "filename": data.get("filename", ""),
            "mimeType": data.get("mimeType", ""),
            "hasBinary": data.get("hasBinary", False),
            "fileSize": data.get("fileSize", 0),
            "parentMessage": data.get("parentMessage", {}),
            "metadata": data.get("metadata", {})
        }
        
        extracted_content = ""
        processing_results = []
        
        # Process binary files if available (from N8N multipart or binary data)
        if files:
            print("📎 Processing uploaded files...")
            for file_key, file in files.items():
                if file and file.filename:
                    binary_data = file.read()
                    result = doc_processor.process_binary_file(
                        binary_data, 
                        file.filename, 
                        file.content_type or "application/octet-stream"
                    )
                    processing_results.append(result)
                    if result.get("success") and result.get("text"):
                        extracted_content += f"\n\nFile: {file.filename}\n{result['text']}"
        
        # Handle N8N binary data format
        elif document_data.get("hasBinary") and "binary" in request.get_json():
            print("📎 Processing N8N binary data...")
            binary_section = request.get_json().get("binary", {})
            
            for key, binary_info in binary_section.items():
                if isinstance(binary_info, dict) and "data" in binary_info:
                    try:
                        # Decode base64 binary data
                        binary_data = base64.b64decode(binary_info["data"])
                        filename = binary_info.get("fileName", binary_info.get("filename", f"attachment_{key}"))
                        mime_type = binary_info.get("mimeType", "application/octet-stream")
                        
                        result = doc_processor.process_binary_file(binary_data, filename, mime_type)
                        processing_results.append(result)
                        
                        if result.get("success") and result.get("text"):
                            extracted_content += f"\n\nFile: {filename}\n{result['text']}"
                            
                    except Exception as e:
                        print(f"❌ Error processing binary data {key}: {e}")
        
        # Process main content (email body/subject) - handle N8N format
        main_content = ""
        
        # Try different content sources from N8N
        if document_data.get('content'):
            main_content += document_data['content']
        
        if document_data.get('subject'):
            main_content += f" {document_data['subject']}"
            
        # Check parentMessage for email content (N8N format)
        if document_data.get('parentMessage'):
            parent_msg = document_data['parentMessage']
            if isinstance(parent_msg, dict):
                if parent_msg.get('subject'):
                    main_content += f" {parent_msg['subject']}"
                if parent_msg.get('content'):
                    main_content += f" {parent_msg['content']}"
                if parent_msg.get('body'):
                    main_content += f" {parent_msg['body']}"
        
        print(f"🔍 Main content extracted: {len(main_content)} characters")
        print(f"🔍 Main content preview: {main_content[:300]}...")
        
        if main_content.strip():
            extracted_content = main_content + extracted_content
        
        # Combine all extracted content for analysis
        content_to_analyze = extracted_content.strip()
        
        if not content_to_analyze:
            return jsonify({"error": "No content to process - no text found in attachments or body"}), 400
        
        print(f"📄 Content to analyze: {len(content_to_analyze)} characters")
        print(f"🔍 Content preview: {content_to_analyze[:200]}...")
        
        # Extract text and detect language
        extraction_result = doc_processor.extract_text_from_content(content_to_analyze)
        
        if "error" in extraction_result:
            return jsonify({"error": extraction_result["error"]}), 500
        
        # Use RAG-enhanced routing if available, fallback to smart router
        if rag_classifier:
            print("🤖 Using RAG classifier for document analysis...")
            
            # First check if this is a multi-department document
            is_multi_dept = False
            multi_dept_result = None
            
            if multi_dept_classifier:
                print("🎯 Checking for multi-department requirements...")
                try:
                    multi_dept_result = multi_dept_classifier.analyze_multi_department_document(content_to_analyze)
                    is_multi_dept = multi_dept_result.get('is_multi_department', False)
                    
                    print(f"📊 Multi-department analysis: {multi_dept_result.get('departments_count', 0)} departments detected")
                    if is_multi_dept:
                        print(f"🏢 Departments: {', '.join(multi_dept_result.get('departments_detected', []))}")
                except Exception as e:
                    print(f"⚠️ Multi-department analysis failed: {e}")
                    is_multi_dept = False
            
            # Get standard RAG analysis
            rag_result = rag_classifier.analyze_with_rag(
                content_to_analyze, 
                document_data.get("subject", "document"), 
                document_data.get("subject", "")
            )
            
            # Enhance routing result with multi-department data if applicable
            if is_multi_dept and multi_dept_result:
                routing_result = {
                    "assigned_department": multi_dept_result.get("primary_department", "Operations"),
                    "priority": rag_result.get("priority", "high"),  # Multi-dept usually high priority
                    "confidence": multi_dept_result.get("confidence", 80),
                    "reasoning": f"Multi-department document requiring coordination across {multi_dept_result.get('departments_count')} departments",
                    "matched_keywords": rag_result.get("primary_keywords", []),
                    "document_type": "multi_department_coordination",
                    "recommended_actions": rag_result.get("recommended_actions", []),
                    "rag_enhanced": True,
                    "multi_department": True,
                    "departments_detected": multi_dept_result.get("departments_detected", []),
                    "department_tasks": multi_dept_result.get("department_specific_tasks", {}),
                    "routing_strategy": multi_dept_result.get("routing_strategy", "multi_department")
                }
                
                # Enhanced AI analysis for multi-department
                ai_result = {
                    "department": multi_dept_result.get("primary_department"),
                    "confidence": multi_dept_result.get("confidence", 80),
                    "summary": multi_dept_result.get("overall_analysis", {}).get("summary", "Multi-department coordination required"),
                    "key_topics": rag_result.get("primary_keywords", []),
                    "priority": "high",
                    "rag_analysis": True,
                    "multi_department_analysis": multi_dept_result
                }
            else:
                # Standard single-department routing
                routing_result = {
                    "assigned_department": rag_result.get("department", "Administration"),
                    "priority": rag_result.get("priority", "normal"),
                    "confidence": rag_result.get("confidence_score", 70),
                    "reasoning": rag_result.get("reasoning", "RAG-enhanced analysis"),
                    "matched_keywords": rag_result.get("primary_keywords", []),
                    "document_type": rag_result.get("document_type", "general"),
                    "recommended_actions": rag_result.get("recommended_actions", []),
                    "rag_enhanced": True,
                    "multi_department": False
                }
                
                # Use RAG result as AI analysis too
                ai_result = {
                    "department": rag_result.get("department"),
                    "confidence": rag_result.get("confidence_score", 70),
                    "summary": rag_result.get("reasoning", ""),
                    "key_topics": rag_result.get("primary_keywords", []),
                    "priority": rag_result.get("priority", "normal"),
                    "rag_analysis": True
                }
        else:
            print("⚠️ RAG classifier not available, using fallback routing...")
            # Fallback to original routing
            routing_result = smart_router.route_document(content_to_analyze, document_data)
            ai_result = ai_analyzer.analyze_document(content_to_analyze[:500])
        
        # Create document record with intelligent duplicate detection
        timestamp = datetime.now()
        
        # Generate content hash for duplicate detection
        content_hash = hashlib.md5(content_to_analyze.encode('utf-8')).hexdigest()
        
        # Check for existing documents with same content hash or message ID
        duplicate_doc = None
        try:
            # First check by content hash
            existing_content = supabase.table("documents").select("*").eq('metadata->>content_hash', content_hash).execute()
            if existing_content.data:
                duplicate_doc = existing_content.data[0]
                logger.info(f"📋 Found duplicate by content hash: {duplicate_doc['id']}")
            
            # If no content match, check by message ID (for same source)
            if not duplicate_doc and document_data.get('message_id'):
                existing_msg = supabase.table("documents").select("*").eq('metadata->>message_id', document_data['message_id']).eq('source', document_data['source']).execute()
                if existing_msg.data:
                    duplicate_doc = existing_msg.data[0]
                    logger.info(f"📨 Found duplicate by message ID: {duplicate_doc['id']}")
                    
        except Exception as e:
            logger.warning(f"⚠️ Error checking for duplicates: {e}")
        
        # If duplicate found, update instead of insert
        if duplicate_doc:
            logger.info(f"🔄 Updating existing document: {duplicate_doc['id']}")
            
            # Update the existing document with new analysis
            update_data = {
                "assigned_department": routing_result["assigned_department"],
                "priority": routing_result["priority"],
                "confidence": routing_result["confidence"],
                "updated_at": timestamp.isoformat(),
                "metadata": {
                    **duplicate_doc.get('metadata', {}),
                    "content_hash": content_hash,
                    "last_processed": timestamp.isoformat(),
                    "routing_reasoning": routing_result["reasoning"],
                    "matched_keywords": routing_result.get("matched_keywords", []),
                    "ai_analysis": ai_result if "error" not in ai_result else None,
                    "rag_enhanced": routing_result.get("rag_enhanced", False),
                    "recommended_actions": routing_result.get("recommended_actions", []),
                    "document_type": routing_result.get("document_type", "general"),
                    "multi_department": routing_result.get("multi_department", False),
                    "departments_detected": routing_result.get("departments_detected", []),
                    "department_tasks": routing_result.get("department_tasks", {}),
                    "routing_strategy": routing_result.get("routing_strategy", "single_department"),
                    "duplicate_prevention": True,
                    "update_count": duplicate_doc.get('metadata', {}).get('update_count', 0) + 1
                }
            }
            
            try:
                result = supabase.table("documents").update(update_data).eq('id', duplicate_doc['id']).execute()
                stored_doc = result.data[0] if result.data else {**duplicate_doc, **update_data}
                logger.info(f"✅ Document updated successfully: {stored_doc['id']}")
            except Exception as e:
                logger.error(f"❌ Error updating document: {e}")
                stored_doc = duplicate_doc
        else:
            # No duplicate found, create new document
            unique_suffix = str(uuid.uuid4())[:8]
            message_id_part = document_data['message_id'][:8] if document_data['message_id'] else "unknown"
            
            processed_doc = {
                "id": f"kmrl_{timestamp.strftime('%Y%m%d_%H%M%S')}_{message_id_part}_{unique_suffix}",
                "title": document_data["subject"] or "KMRL Document",
                "content": content_to_analyze,
                "source": document_data["source"],
                "language": extraction_result["language"],
                "assigned_department": routing_result["assigned_department"],
                "priority": routing_result["priority"],
                "confidence": routing_result["confidence"],
                "status": "pending",
                "created_at": timestamp.isoformat(),
                "metadata": {
                    "content_hash": content_hash,
                    "sender": document_data["sender"],
                    "original_subject": document_data["subject"],
                    "timestamp": document_data["timestamp"],
                    "message_id": document_data["message_id"],
                    "word_count": extraction_result["word_count"],
                    "routing_reasoning": routing_result["reasoning"],
                    "matched_keywords": routing_result.get("matched_keywords", []),
                    "ai_analysis": ai_result if "error" not in ai_result else None,
                    "attachments_count": len(document_data["attachments"]),
                    "rag_enhanced": routing_result.get("rag_enhanced", False),
                    "recommended_actions": routing_result.get("recommended_actions", []),
                    "document_type": routing_result.get("document_type", "general"),
                    "multi_department": routing_result.get("multi_department", False),
                    "departments_detected": routing_result.get("departments_detected", []),
                    "department_tasks": routing_result.get("department_tasks", {}),
                    "routing_strategy": routing_result.get("routing_strategy", "single_department"),
                    "duplicate_prevention": True,
                    "update_count": 0
                }
            }
            
            # Store new document in Supabase
            try:
                result = supabase.table("documents").insert(processed_doc).execute()
                stored_doc = result.data[0] if result.data else processed_doc
                logger.info(f"✅ New document stored successfully: {stored_doc['id']}")
            except Exception as e:
                error_message = str(e)
                logger.error(f"❌ Error storing new document: {error_message}")
                
                # If still getting duplicate errors, generate completely new ID
                if "duplicate key" in error_message.lower() or "23505" in error_message:
                    logger.info("🔄 Primary key conflict, generating ultra-unique ID...")
                    processed_doc["id"] = f"kmrl_{timestamp.strftime('%Y%m%d_%H%M%S_%f')}_{str(uuid.uuid4())[:12]}"
                    
                    try:
                        result = supabase.table("documents").insert(processed_doc).execute()
                        stored_doc = result.data[0] if result.data else processed_doc
                        logger.info(f"✅ Document stored with ultra-unique ID: {stored_doc['id']}")
                    except Exception as retry_error:
                        logger.error(f"❌ Critical: Failed to store document: {str(retry_error)}")
                        stored_doc = processed_doc
                else:
                    stored_doc = processed_doc
        
        response_data = {
            "status": "success",
            "message": "Document updated successfully" if duplicate_doc else "Document processed successfully",
            "document": {
                "id": stored_doc["id"],
                "title": stored_doc["title"],
                "assigned_department": stored_doc["assigned_department"],
                "priority": stored_doc["priority"],
                "confidence": stored_doc["confidence"],
                "language": stored_doc["language"],
                "reasoning": routing_result["reasoning"],
                "is_update": bool(duplicate_doc)
            }
        }
        
        # Add multi-department information if applicable
        if routing_result.get("multi_department", False):
            response_data["document"]["multi_department"] = True
            response_data["document"]["departments_detected"] = routing_result.get("departments_detected", [])
            response_data["document"]["departments_count"] = len(routing_result.get("departments_detected", []))
            response_data["document"]["routing_strategy"] = routing_result.get("routing_strategy", "multi_department")
            response_data["message"] = f"Multi-department document processed - requires coordination across {len(routing_result.get('departments_detected', []))} departments"
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get documents by department, status, or all"""
    try:
        department = request.args.get('department')
        status = request.args.get('status')
        priority = request.args.get('priority')
        limit = int(request.args.get('limit', 50))
        
        query = supabase.table("documents").select("*")
        
        if department:
            query = query.eq('assigned_department', department)
        
        if status:
            query = query.eq('status', status)
            
        if priority:
            query = query.eq('priority', priority)
        
        result = query.order('created_at', desc=True).limit(limit).execute()
        
        return jsonify({
            "status": "success",
            "documents": result.data,
            "count": len(result.data),
            "filters": {
                "department": department,
                "status": status,
                "priority": priority,
                "limit": limit
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/documents/<document_id>', methods=['GET'])
def get_single_document(document_id):
    """Get a single document by ID"""
    try:
        # Check if request wants HTML view
        if request.args.get('view') == 'html':
            return view_document_html(document_id)
        
        # Get document from database
        result = supabase.table("documents").select("*").eq("id", document_id).execute()
        
        if not result.data:
            return jsonify({"error": "Document not found"}), 404
            
        document = result.data[0]
        
        return jsonify({
            "status": "success",
            "document": document
        })
        
    except Exception as e:
        logger.error(f"Error fetching document {document_id}: {str(e)}")
        return jsonify({"error": str(e)}), 500

def view_document_html(document_id):
    """Serve document as HTML for viewing"""
    try:
        result = supabase.table("documents").select("*").eq("id", document_id).execute()
        
        if not result.data:
            return "<h1>Document not found</h1>", 404
            
        document = result.data[0]
        content = document.get('content', '')
        title = document.get('title', 'Untitled Document')
        source = document.get('source', 'Unknown')
        department = document.get('assigned_department', 'Unknown')
        
        # Format content for HTML display
        formatted_content = content.replace('\n', '<br>')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{title}</title>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
                .header {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                .content {{ background: white; padding: 20px; border: 1px solid #ddd; border-radius: 5px; }}
                .meta {{ color: #666; font-size: 0.9em; margin-bottom: 10px; }}
                .title {{ color: #333; margin-bottom: 10px; }}
                .department {{ background: #e3f2fd; color: #1976d2; padding: 5px 10px; border-radius: 3px; display: inline-block; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1 class="title">{title}</h1>
                <div class="meta">
                    <span class="department">{department}</span>
                    <span style="margin-left: 15px;">Source: {source}</span>
                </div>
            </div>
            <div class="content">
                {formatted_content}
            </div>
        </body>
        </html>
        """
        
        return html_content, 200, {'Content-Type': 'text/html'}
        
    except Exception as e:
        return f"<h1>Error loading document</h1><p>{str(e)}</p>", 500

@app.route('/api/analysis/document/<document_id>', methods=['GET'])
def get_individual_document(document_id):
    """Get individual document with enhanced analysis for frontend"""
    try:
        # Get document from database
        result = supabase.table("documents").select("*").eq("id", document_id).execute()
        
        if not result.data:
            return jsonify({"error": "Document not found"}), 404
            
        doc = result.data[0]
        
        # Transform document for frontend format
        transformed_doc = {
            "id": doc.get("id"),
            "title": doc.get("title"),
            "content": doc.get("content"),
            "source": doc.get("source"),
            "format": doc.get("mime_type", "text/plain"),
            "assigned_department": doc.get("assigned_department"),
            "confidence": doc.get("confidence", 0),
            "priority": doc.get("priority", "normal"),
            "reasoning": "",
            "language": doc.get("language", "english"),
            "metadata": {
                "file_info": {
                    "filename": doc.get("title", "unknown"),
                    "mimeType": doc.get("mime_type", "text/plain"),
                    "size": doc.get("file_size", 0),
                    "hasDownload": False,
                    "hasBinary": doc.get("binary_data") is not None
                },
                "original_data": {
                    "parentMessage": doc.get("metadata", {}).get("parent_message", {})
                },
                "processing_timestamp": doc.get("created_at")
            },
            "department_details": {
                "code": doc.get("assigned_department", "").upper()[:3],
                "malayalam_name": "",
                "head": "",
                "contact_info": {
                    "email": "",
                    "phone": ""
                },
                "due_date": "",
                "expected_response": ""
            },
            "advanced_analysis": {
                "ai_insights": {
                    "summary": "",
                    "key_topics": [],
                    "recommended_actions": [],
                    "action_required": False,
                    "confidence_score": doc.get("confidence", 0)
                }
            }
        }
        
        # Extract enhanced analysis from metadata if available
        metadata = doc.get("metadata", {})
        ai_analysis = metadata.get("ai_analysis", {})
        
        if ai_analysis:
            # Update with enhanced RAG analysis
            transformed_doc["advanced_analysis"]["ai_insights"].update({
                "summary": ai_analysis.get("document_summary", ""),
                "key_topics": ai_analysis.get("main_topics", []),
                "recommended_actions": ai_analysis.get("recommended_actions", []),
                "action_required": ai_analysis.get("priority", "normal") in ["urgent", "high"],
                "confidence_score": ai_analysis.get("confidence_score", doc.get("confidence", 0))
            })
            transformed_doc["reasoning"] = ai_analysis.get("reasoning", "")
        
        return jsonify({
            "status": "success",
            "document": transformed_doc
        })
        
    except Exception as e:
        logger.error(f"Error fetching document {document_id}: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/departments/stats', methods=['GET'])
def get_department_stats():
    """Get statistics by department"""
    try:
        # Get document counts by department
        result = supabase.table("documents").select("assigned_department, priority, status").execute()
        
        stats = {}
        for doc in result.data:
            dept = doc["assigned_department"]
            if dept not in stats:
                stats[dept] = {
                    "total": 0,
                    "pending": 0,
                    "in_progress": 0,
                    "completed": 0,
                    "urgent": 0,
                    "high": 0,
                    "normal": 0
                }
            
            stats[dept]["total"] += 1
            stats[dept][doc["status"]] = stats[dept].get(doc["status"], 0) + 1
            stats[dept][doc["priority"]] = stats[dept].get(doc["priority"], 0) + 1
        
        return jsonify({
            "status": "success",
            "department_stats": stats
        })
        
    except Exception as e:
        logger.error(f"Error fetching department stats: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/departments/<department>/tasks', methods=['GET'])
def get_department_tasks(department):
    """Get department-specific tasks including from multi-department documents"""
    try:
        # Get documents assigned to this department
        primary_docs_result = supabase.table("documents").select("*").eq('assigned_department', department).order('created_at', desc=True).limit(50).execute()
        
        # Get multi-department documents that have tasks for this department
        all_docs_result = supabase.table("documents").select("*").order('created_at', desc=True).limit(100).execute()
        
        primary_documents = primary_docs_result.data
        multi_dept_tasks = []
        
        # Process multi-department documents
        for doc in all_docs_result.data:
            metadata = doc.get('metadata', {})
            
            # Check if this is a multi-department document
            if metadata.get('multi_department', False):
                departments_detected = metadata.get('departments_detected', [])
                department_tasks = metadata.get('department_tasks', {})
                
                # Normalize department name for matching
                dept_variations = {
                    'Engineering': ['engineering', 'civil', 'mechanical', 'track'],
                    'Operations': ['operations', 'station operations', 'facilities'],
                    'Safety & Security': ['safety', 'security', 'safety & security'],
                    'Electrical': ['electrical', 'power', 'electrical & power'],
                    'Rolling Stock & Mechanical': ['rolling stock', 'mechanical', 'rolling stock & mechanical'],
                    'Signalling': ['signalling', 'communication', 'signalling & communication'],
                    'Environment': ['environment', 'cleaning', 'facilities management']
                }
                
                # Check if this department is involved
                department_matched = False
                matched_dept_name = department
                
                # Direct match
                if department in departments_detected:
                    department_matched = True
                    matched_dept_name = department
                # Check variations
                else:
                    for standard_dept, variations in dept_variations.items():
                        if (department.lower() in [v.lower() for v in variations] and 
                            standard_dept in departments_detected):
                            department_matched = True
                            matched_dept_name = standard_dept
                            break
                        elif (standard_dept.lower() == department.lower() and 
                              any(var in departments_detected for var in variations)):
                            department_matched = True
                            matched_dept_name = standard_dept
                            break
                
                if department_matched:
                    # Get tasks for this department
                    dept_tasks = department_tasks.get(matched_dept_name, [])
                    if not dept_tasks:
                        # Try other variations
                        for standard_dept, variations in dept_variations.items():
                            if department.lower() in [v.lower() for v in variations]:
                                dept_tasks = department_tasks.get(standard_dept, [])
                                if dept_tasks:
                                    break
                    
                    # Get other departments (excluding this one)
                    other_departments = [d for d in departments_detected if d != matched_dept_name]
                    
                    if dept_tasks:  # Only include if there are actual tasks
                        multi_dept_tasks.append({
                            'document': doc,
                            'tasks': dept_tasks,
                            'otherDepartments': other_departments
                        })
        
        return jsonify({
            "status": "success",
            "documents": primary_documents,
            "multiDepartmentTasks": multi_dept_tasks,
            "department": department,
            "summary": {
                "primary_documents": len(primary_documents),
                "multi_department_tasks": len(multi_dept_tasks),
                "total_tasks": len(primary_documents) + len(multi_dept_tasks)
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching department tasks for {department}: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/departments/<department>/tasks/<document_id>', methods=['POST'])
def update_task_status(department, document_id):
    """Update task completion status for a department-document combination"""
    try:
        data = request.get_json()
        task_index = data.get('task_index', 0)
        completed = data.get('completed', False)
        notes = data.get('notes', '')
        
        # Create or update task status record
        task_status = {
            'document_id': document_id,
            'department': department,
            'task_index': task_index,
            'completed': completed,
            'completed_at': datetime.now().isoformat() if completed else None,
            'notes': notes,
            'updated_at': datetime.now().isoformat()
        }
        
        # Try to update existing record first
        existing_result = supabase.table('task_status').select('*').eq('document_id', document_id).eq('department', department).eq('task_index', task_index).execute()
        
        if existing_result.data:
            # Update existing
            result = supabase.table('task_status').update(task_status).eq('document_id', document_id).eq('department', department).eq('task_index', task_index).execute()
        else:
            # Insert new
            task_status['created_at'] = datetime.now().isoformat()
            result = supabase.table('task_status').insert(task_status).execute()
        
        return jsonify({
            "status": "success",
            "message": f"Task {'completed' if completed else 'updated'}",
            "task_status": task_status
        })
        
    except Exception as e:
        logger.error(f"Error updating task status: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/departments/<department>/tasks/<document_id>', methods=['GET'])
def get_task_statuses(department, document_id):
    """Get task completion statuses for a department-document combination"""
    try:
        result = supabase.table('task_status').select('*').eq('document_id', document_id).eq('department', department).execute()
        
        return jsonify({
            "status": "success",
            "task_statuses": result.data
        })
        
    except Exception as e:
        logger.error(f"Error fetching task statuses: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/chat', methods=['POST'])
def chat_with_documents():
    """Enhanced chat interface for intelligent document interaction"""
    try:
        data = request.get_json()
        query = data.get("query", "")
        department = data.get("department")
        user_id = data.get("user_id", "anonymous")  # For privacy management
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        query_lower = query.lower()
        
        # Handle different types of queries
        
        # 1. Document Summary Requests
        if "summary" in query_lower or "what is this document" in query_lower or "document about" in query_lower:
            return handle_document_summary_request(query, department)
        
        # 2. Upload with Duplicate Check
        elif "upload" in query_lower and ("globally" in query_lower or "privately" in query_lower):
            return handle_smart_upload_request(query, department, user_id)
        
        # 3. Document Analysis Questions
        elif any(word in query_lower for word in ["analyze", "what does", "explain", "details", "content"]):
            return handle_document_analysis_request(query, department, user_id)
        
        # 4. General Document Search
        else:
            return handle_general_document_search(query, department, user_id)
        
    except Exception as e:
        logger.error(f"Error in chat: {str(e)}")
        return jsonify({"error": str(e)}), 500


def handle_document_summary_request(query, department):
    """Handle requests for document summaries"""
    try:
        # Get recent documents for the department
        doc_query = supabase.table("documents").select("*")
        if department:
            doc_query = doc_query.eq('assigned_department', department)
        
        documents = doc_query.order('created_at', desc=True).limit(5).execute().data
        
        if not documents:
            return jsonify({
                "answer": f"No documents found for {department} department.",
                "type": "summary",
                "documents": []
            })
        
        # Generate summaries using OpenRouter
        summaries = []
        for doc in documents:
            try:
                if multi_dept_classifier:
                    # Use AI to generate intelligent summary
                    summary_prompt = f"Provide a brief summary of this document:\\n\\nTitle: {doc['title']}\\nContent: {doc['content'][:500]}..."
                    summary = multi_dept_classifier.analyze_document(summary_prompt, {'task': 'summarize'})
                    summary_text = summary.get('summary', f"Document: {doc['title']} - Contains {len(doc['content'])} characters of content")
                else:
                    # Fallback summary
                    summary_text = f"Document: {doc['title']} - {doc['content'][:100]}..."
                
                summaries.append({
                    "title": doc['title'],
                    "summary": summary_text,
                    "department": doc['assigned_department'],
                    "priority": doc['priority'],
                    "id": doc['id']
                })
            except:
                summaries.append({
                    "title": doc['title'],
                    "summary": f"Document contains {len(doc['content'])} characters of content",
                    "department": doc['assigned_department'],
                    "priority": doc['priority'],
                    "id": doc['id']
                })
        
        response = f"Here are summaries of recent documents for {department}:\\n"
        for i, summary in enumerate(summaries, 1):
            response += f"\\n{i}. **{summary['title']}**\\n   {summary['summary']}\\n   Priority: {summary['priority']}"
        
        return jsonify({
            "answer": response,
            "type": "summary",
            "documents": summaries
        })
        
    except Exception as e:
        return jsonify({"error": f"Summary generation failed: {str(e)}"}), 500


def handle_smart_upload_request(query, department, user_id):
    """Handle upload requests with duplicate detection and privacy options"""
    try:
        is_private = "privately" in query.lower()
        is_global = "globally" in query.lower()
        
        response = {
            "answer": f"To upload a document {'privately' if is_private else 'globally'}, please use the file upload interface.\\n\\n",
            "type": "upload_instruction",
            "upload_options": {
                "privacy": "private" if is_private else "global",
                "duplicate_check": True,
                "instructions": []
            }
        }
        
        if is_private:
            response["answer"] += "📁 **Private Upload**: Document will only be visible to you and won't appear in global search.\\n"
            response["upload_options"]["instructions"].append("Document will be stored privately")
        else:
            response["answer"] += "🌐 **Global Upload**: Document will be visible to all departments and appear in global search.\\n"
            response["upload_options"]["instructions"].append("Document will be stored globally")
        
        response["answer"] += "\\n✅ **Duplicate Check**: System will automatically check for existing documents with same content.\\n"
        response["answer"] += "\\n🤖 **Smart Routing**: AI will determine the best department for your document."
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({"error": f"Upload instruction failed: {str(e)}"}), 500


def handle_document_analysis_request(query, department, user_id):
    """Handle document analysis and Q&A requests"""
    try:
        # Search for relevant documents based on query
        doc_query = supabase.table("documents").select("*")
        
        # Include private documents for the user
        if department:
            public_docs = doc_query.eq('assigned_department', department).is_('private_to_user', 'null').execute().data
            private_docs = doc_query.eq('assigned_department', department).eq('private_to_user', user_id).execute().data
            documents = public_docs + private_docs
        else:
            public_docs = supabase.table("documents").select("*").is_('private_to_user', 'null').execute().data
            private_docs = supabase.table("documents").select("*").eq('private_to_user', user_id).execute().data
            documents = public_docs + private_docs
        
        # Find relevant documents using keyword matching
        query_words = query.lower().split()
        relevant_docs = []
        
        for doc in documents:
            content_lower = doc["content"].lower()
            title_lower = doc["title"].lower()
            
            # Calculate relevance score
            score = 0
            for word in query_words:
                if word in title_lower:
                    score += 3
                if word in content_lower:
                    score += 1
            
            if score > 0:
                relevant_docs.append((doc, score))
        
        # Sort by relevance
        relevant_docs.sort(key=lambda x: x[1], reverse=True)
        
        if not relevant_docs:
            return jsonify({
                "answer": "I couldn't find any documents related to your question. Please try a different query or upload a relevant document.",
                "type": "analysis",
                "sources": []
            })
        
        # Generate intelligent response using the most relevant documents
        top_docs = [doc for doc, score in relevant_docs[:3]]
        
        try:
            if multi_dept_classifier:
                # Create context from relevant documents
                context = "\\n\\n".join([f"Document: {doc['title']}\\nContent: {doc['content'][:300]}..." for doc in top_docs])
                analysis_prompt = f"Based on these documents, answer this question: {query}\\n\\nContext:\\n{context}"
                
                analysis = multi_dept_classifier.analyze_document(analysis_prompt, {'task': 'question_answering'})
                answer = analysis.get('summary', f"Based on {len(top_docs)} relevant documents, here's what I found...")
            else:
                # Fallback response
                answer = f"Based on {len(top_docs)} relevant documents:\\n"
                for doc in top_docs:
                    answer += f"\\n• {doc['title']}: {doc['content'][:100]}..."
        
        except Exception as ai_error:
            logger.warning(f"AI analysis failed: {ai_error}")
            answer = f"I found {len(top_docs)} relevant documents:\\n"
            for doc in top_docs:
                answer += f"\\n• **{doc['title']}** ({doc['assigned_department']})\\n  {doc['content'][:150]}..."
        
        return jsonify({
            "answer": answer,
            "type": "analysis",
            "sources": [{"title": doc["title"], "id": doc["id"], "department": doc["assigned_department"]} for doc in top_docs],
            "relevant_documents": len(relevant_docs)
        })
        
    except Exception as e:
        return jsonify({"error": f"Document analysis failed: {str(e)}"}), 500


def handle_general_document_search(query, department, user_id):
    """Handle general document search queries"""
    try:
        # Search in both public and user's private documents
        doc_query = supabase.table("documents").select("*")
        
        if department:
            public_docs = doc_query.eq('assigned_department', department).is_('private_to_user', 'null').execute().data
            private_docs = doc_query.eq('assigned_department', department).eq('private_to_user', user_id).execute().data
            documents = public_docs + private_docs
        else:
            public_docs = supabase.table("documents").select("*").is_('private_to_user', 'null').limit(20).execute().data
            private_docs = supabase.table("documents").select("*").eq('private_to_user', user_id).limit(10).execute().data
            documents = public_docs + private_docs
        
        # Simple keyword matching
        query_lower = query.lower()
        relevant_docs = [
            doc for doc in documents 
            if any(word in doc["content"].lower() or word in doc["title"].lower() for word in query_lower.split())
        ]
        
        if relevant_docs:
            response_text = f"Found {len(relevant_docs)} relevant documents:\\n"
            for i, doc in enumerate(relevant_docs[:5], 1):
                privacy_indicator = "🔒" if doc.get('private_to_user') else "🌐"
                response_text += f"\\n{i}. {privacy_indicator} **{doc['title']}** ({doc['assigned_department']}, {doc['priority']} priority)"
        else:
            response_text = "No relevant documents found. Try different keywords or upload a relevant document."
        
        return jsonify({
            "answer": response_text,
            "type": "search",
            "relevant_documents": len(relevant_docs),
            "sources": [{"title": doc["title"], "id": doc["id"], "private": bool(doc.get('private_to_user'))} for doc in relevant_docs[:5]]
        })
        
    except Exception as e:
        return jsonify({"error": f"Search failed: {str(e)}"}), 500


@app.route('/api/chat/upload', methods=['POST'])
def chat_upload_with_prompt():
    """Handle file uploads with user prompts in chat interface"""
    try:
        # Check if file is provided
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        user_prompt = request.form.get('prompt', '')
        department = request.form.get('department', 'General')
        user_id = request.form.get('user_id', 'anonymous')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not user_prompt.strip():
            return jsonify({"error": "Please provide a prompt explaining what you want to do with this document"}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        file.save(temp_path)
        
        try:
            # Extract content from file
            content = ""
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                content = extract_text_from_pdf(temp_path)
            elif file_ext in ['doc', 'docx']:
                content = extract_text_from_word(temp_path)
            elif file_ext in ['xls', 'xlsx']:
                content = extract_text_from_excel(temp_path)
            elif file_ext in ['txt']:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_ext in ['csv']:
                import pandas as pd
                df = pd.read_csv(temp_path)
                content = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\\n{df.head(10).to_string()}"
            else:
                content = f"File: {filename} (Type: {file_ext})"
            
            # Analyze user intent from prompt
            intent_analysis = analyze_user_intent(user_prompt, filename, content)
            
            # Generate intelligent response based on intent
            if intent_analysis['action'] == 'upload':
                response = handle_chat_upload_action(filename, content, user_prompt, intent_analysis, department, user_id)
            elif intent_analysis['action'] == 'analyze':
                response = handle_chat_analyze_action(filename, content, user_prompt, intent_analysis)
            elif intent_analysis['action'] == 'summarize':
                response = handle_chat_summarize_action(filename, content, user_prompt, intent_analysis)
            elif intent_analysis['action'] == 'question':
                response = handle_chat_question_action(filename, content, user_prompt, intent_analysis, department, user_id)
            else:
                response = handle_chat_general_action(filename, content, user_prompt, intent_analysis)
            
            return jsonify(response)
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        logger.error(f"❌ Error in chat upload: {e}")
        return jsonify({"error": str(e)}), 500


def analyze_user_intent(prompt: str, filename: str, content: str) -> Dict[str, Any]:
    """Analyze user intent from their prompt"""
    prompt_lower = prompt.lower()
    
    # Intent classification
    if any(word in prompt_lower for word in ['upload', 'store', 'save', 'add to database', 'keep this']):
        action = 'upload'
        if any(word in prompt_lower for word in ['private', 'personal', 'only me', 'confidential']):
            privacy = 'private'
        else:
            privacy = 'global'
    elif any(word in prompt_lower for word in ['analyze', 'examine', 'study', 'review', 'check']):
        action = 'analyze'
        privacy = 'none'
    elif any(word in prompt_lower for word in ['summarize', 'summary', 'brief', 'overview', 'what is this']):
        action = 'summarize'
        privacy = 'none'
    elif any(word in prompt_lower for word in ['?', 'how', 'what', 'why', 'when', 'where', 'explain']):
        action = 'question'
        privacy = 'none'
    else:
        action = 'general'
        privacy = 'none'
    
    # Extract specific requests
    urgency = 'normal'
    if any(word in prompt_lower for word in ['urgent', 'asap', 'immediately', 'priority']):
        urgency = 'high'
    
    return {
        'action': action,
        'privacy': privacy,
        'urgency': urgency,
        'raw_prompt': prompt,
        'filename': filename,
        'content_preview': content[:200] + '...' if len(content) > 200 else content
    }


def handle_chat_upload_action(filename: str, content: str, prompt: str, intent: Dict, department: str, user_id: str) -> Dict[str, Any]:
    """Handle upload action from chat"""
    try:
        # Check for duplicates
        duplicate_check = check_document_exists(filename, content, user_id)
        
        if duplicate_check['is_duplicate']:
            existing_doc = duplicate_check['existing_content_docs'][0]
            return {
                "answer": f"🔍 **Duplicate Detected!**\\n\\nA document with identical content already exists:\\n• **{existing_doc['title']}** (uploaded {existing_doc.get('created_at', 'previously')})\\n• Department: {existing_doc.get('assigned_department')}\\n\\n❓ Would you like me to upload it anyway as a new version?",
                "type": "duplicate_warning",
                "duplicate_info": {
                    "existing_title": existing_doc['title'],
                    "existing_id": existing_doc['id'],
                    "existing_department": existing_doc.get('assigned_department')
                },
                "upload_anyway": True
            }
        
        # Determine department using AI
        if multi_dept_classifier:
            routing_result = multi_dept_classifier.analyze_document(content, {'filename': filename, 'user_intent': prompt})
            assigned_department = routing_result.get('departments', [department])[0] if routing_result else department
        else:
            assigned_department = department
        
        # Store document
        content_hash = calculate_content_hash(content)
        doc_data = {
            'title': filename,
            'content': content,
            'source': f"Chat Upload - {assigned_department}",
            'assigned_department': assigned_department,
            'status': 'processed',
            'priority': intent['urgency'],
            'private_to_user': user_id if intent['privacy'] == 'private' else None,
            'metadata': {
                'uploaded_via': 'chat_interface',
                'user_prompt': prompt,
                'content_hash': content_hash,  # Store in metadata
                'privacy_mode': intent['privacy'],
                'user_intent': intent,
                'ai_routed_department': assigned_department,
                'ai_analysis': {
                    'intent': intent,
                    'department': assigned_department,
                    'processing_method': 'chat_upload'
                }
            }
        }
        
        result = supabase.table('documents').insert(doc_data).execute()
        
        privacy_text = "privately (only visible to you)" if intent['privacy'] == 'private' else "globally (visible to all departments)"
        
        return {
            "answer": f"✅ **Document Uploaded Successfully!**\\n\\n📁 **File:** {filename}\\n🏢 **Department:** {assigned_department}\\n🔒 **Privacy:** {privacy_text}\\n📝 **Your Intent:** {prompt}\\n\\n🤖 The document has been intelligently routed to the {assigned_department} department and is now available in their dashboard.",
            "type": "upload_success",
            "document_id": result.data[0]['id'],
            "department": assigned_department,
            "privacy": intent['privacy']
        }
        
    except Exception as e:
        return {
            "answer": f"❌ **Upload Failed:** {str(e)}",
            "type": "upload_error"
        }


def handle_chat_analyze_action(filename: str, content: str, prompt: str, intent: Dict) -> Dict[str, Any]:
    """Handle analysis action from chat"""
    try:
        if multi_dept_classifier:
            # Use AI to analyze document based on user's specific request
            analysis_prompt = f"User wants to: {prompt}\\n\\nDocument: {filename}\\nContent: {content[:1000]}..."
            analysis = multi_dept_classifier.analyze_document(analysis_prompt, {'task': 'detailed_analysis', 'user_request': prompt})
            
            ai_response = analysis.get('summary', f"Document Analysis for {filename}")
        else:
            ai_response = f"Document: {filename}\\nContent Length: {len(content)} characters\\nFile appears to contain structured data."
        
        return {
            "answer": f"🔍 **Document Analysis**\\n\\n📄 **File:** {filename}\\n📝 **Your Request:** {prompt}\\n\\n**Analysis Results:**\\n{ai_response}",
            "type": "analysis",
            "filename": filename
        }
        
    except Exception as e:
        return {
            "answer": f"❌ **Analysis Failed:** {str(e)}",
            "type": "analysis_error"
        }


def handle_chat_summarize_action(filename: str, content: str, prompt: str, intent: Dict) -> Dict[str, Any]:
    """Handle summarization action from chat"""
    try:
        # Create summary based on content
        if len(content) > 500:
            summary = content[:500] + "..."
        else:
            summary = content
        
        # Count basic stats
        word_count = len(content.split())
        char_count = len(content)
        
        return {
            "answer": f"📋 **Document Summary**\\n\\n📄 **File:** {filename}\\n📝 **Your Request:** {prompt}\\n\\n**Quick Stats:**\\n• Words: {word_count}\\n• Characters: {char_count}\\n\\n**Content Preview:**\\n{summary}\\n\\n💡 Would you like me to upload this document or perform any other actions?",
            "type": "summary",
            "filename": filename,
            "stats": {"words": word_count, "characters": char_count}
        }
        
    except Exception as e:
        return {
            "answer": f"❌ **Summary Failed:** {str(e)}",
            "type": "summary_error"
        }


def handle_chat_question_action(filename: str, content: str, prompt: str, intent: Dict, department: str, user_id: str) -> Dict[str, Any]:
    """Handle question action from chat"""
    try:
        # Try to answer the question based on document content
        if multi_dept_classifier:
            qa_prompt = f"Question: {prompt}\\n\\nDocument Context:\\nFilename: {filename}\\nContent: {content[:1000]}..."
            qa_result = multi_dept_classifier.analyze_document(qa_prompt, {'task': 'question_answering'})
            answer = qa_result.get('summary', f"Based on the document {filename}, here's what I found...")
        else:
            # Simple keyword matching fallback
            answer = f"In the document '{filename}', I found content related to your question. The document contains {len(content)} characters of information."
        
        return {
            "answer": f"❓ **Question & Answer**\\n\\n📄 **Document:** {filename}\\n❓ **Your Question:** {prompt}\\n\\n**Answer:**\\n{answer}\\n\\n💡 Would you like me to upload this document for future reference?",
            "type": "question_answer",
            "filename": filename
        }
        
    except Exception as e:
        return {
            "answer": f"❌ **Question Processing Failed:** {str(e)}",
            "type": "qa_error"
        }


def handle_chat_general_action(filename: str, content: str, prompt: str, intent: Dict) -> Dict[str, Any]:
    """Handle general/unclear actions from chat"""
    return {
        "answer": f"📄 **Document Received: {filename}**\\n\\n📝 **Your Message:** {prompt}\\n\\n🤖 I can help you with this document in several ways:\\n\\n1. **📤 Upload it** - 'Upload this globally/privately'\\n2. **🔍 Analyze it** - 'Analyze this document'\\n3. **📋 Summarize it** - 'Give me a summary'\\n4. **❓ Answer questions** - Ask specific questions about the content\\n\\nWhat would you like me to do with this document?",
        "type": "general_help",
        "filename": filename,
        "suggestions": ["Upload globally", "Upload privately", "Analyze content", "Summarize", "Ask questions"]
    }

# RAG System Endpoints
@app.route('/api/rag/stats', methods=['GET'])
def get_rag_stats():
    """Get RAG system statistics"""
    try:
        from rag_classifier import RAGDepartmentClassifier
        rag = RAGDepartmentClassifier()
        
        stats = {
            "knowledge_base_size": len(rag.knowledge_base),
            "vector_index_built": rag.document_vectors is not None,
            "vectorizer_fitted": hasattr(rag.vectorizer, 'vocabulary_'),
            "departments_covered": list(set([doc.get('department', 'Unknown') for doc in rag.knowledge_base])),
            "total_documents_in_kb": len(rag.knowledge_base),
            "status": "operational" if rag.knowledge_base else "not_initialized"
        }
        
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"RAG stats error: {e}")
        return jsonify({"error": str(e), "status": "error"}), 500

@app.route('/api/rag/retrain', methods=['POST'])
def retrain_rag():
    """Retrain RAG system with latest documents"""
    try:
        from rag_classifier import RAGDepartmentClassifier
        
        # Get recent documents from database
        result = supabase.table('documents').select('*').limit(100).execute()
        recent_docs = result.data if result.data else []
        
        # Initialize RAG and add documents to knowledge base
        rag = RAGDepartmentClassifier()
        
        added_count = 0
        for doc in recent_docs:
            if doc.get('content') and doc.get('department'):
                rag.learn_from_document({
                    'content': doc['content'],
                    'department': doc['department'],
                    'filename': doc.get('filename', ''),
                    'keywords': doc.get('key_topics', [])
                })
                added_count += 1
        
        # Rebuild vector index
        rag.build_vector_index()
        rag.save_knowledge_base()
        
        return jsonify({
            "message": "RAG system retrained successfully",
            "documents_added": added_count,
            "total_kb_size": len(rag.knowledge_base),
            "timestamp": datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"RAG retrain error: {e}")
        return jsonify({"error": str(e)}), 500

# Import document storage service
try:
    from document_storage import store_document_handler
except ImportError:
    store_document_handler = None
    print("⚠️ Document storage service not available")

def extract_content_from_n8n_data(data, parent_message, metadata):
    """
    Enhanced content extraction from N8N data when binary data is not directly available
    """
    try:
        logger.info(f"🔍 Analyzing N8N data for content extraction")
        
        # Check if there's any embedded content in the data
        potential_content_fields = [
            'content', 'text', 'body', 'html', 'snippet',
            'email_content', 'attachment_content', 'file_content'
        ]
        
        extracted_text = ""
        binary_data = None
        
        # Method 1: Check for any base64 encoded content in the data
        for field in potential_content_fields:
            field_value = data.get(field, '')
            if field_value and len(field_value) > 100:
                # Check if it looks like base64
                if is_base64_content(field_value):
                    try:
                        import base64
                        decoded_data = base64.b64decode(field_value)
                        # Try to extract text from the decoded data
                        filename = data.get('filename', 'document')
                        file_extension = filename.lower().split('.')[-1] if '.' in filename else 'txt'
                        
                        if file_extension == 'csv':
                            extracted_text = extract_text_from_csv(decoded_data)
                        elif file_extension in ['txt', 'text']:
                            extracted_text = decoded_data.decode('utf-8', errors='ignore')
                        else:
                            extracted_text = decoded_data.decode('utf-8', errors='ignore')
                        
                        if extracted_text:
                            binary_data = decoded_data
                            logger.info(f"✅ Extracted content from {field}: {len(extracted_text)} chars")
                            break
                    except:
                        continue
        
        # Method 2: Check parent message for attachment info
        if not extracted_text and parent_message:
            # Sometimes N8N includes attachment info in parent message
            if 'attachment_data' in parent_message or 'attachments' in parent_message:
                logger.info(f"📎 Found attachment info in parent message")
                # Process attachment data if available
        
        # Method 3: Use the existing content but try to enhance it
        if not extracted_text and data.get('content'):
            existing_content = data['content']
            # If the content looks like it might be CSV data embedded as text
            if ',' in existing_content and '\\n' in existing_content:
                # Try to parse as CSV-like content
                csv_content = existing_content.replace('\\n', '\n').replace('\\t', '\t')
                if is_csv_like_content(csv_content):
                    extracted_text = csv_content
                    logger.info(f"✅ Extracted CSV-like content: {len(extracted_text)} chars")
        
        if extracted_text:
            return {
                'text': extracted_text,
                'binary_data': binary_data,
                'source': 'n8n_enhanced_extraction'
            }
        
        return None
        
    except Exception as e:
        logger.error(f"❌ Enhanced content extraction error: {e}")
        return None

def is_base64_content(content):
    """Check if content looks like base64 encoding"""
    try:
        if len(content) < 50:
            return False
        # Check if it's mostly base64 characters
        base64_chars = set('ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=')
        content_chars = set(content.replace(' ', '').replace('\n', '').replace('\r', ''))
        return len(content_chars - base64_chars) / len(content_chars) < 0.1
    except:
        return False

def is_csv_like_content(content):
    """Check if content looks like CSV data"""
    try:
        lines = content.split('\n')
        if len(lines) < 2:
            return False
        
        # Check if first few lines have consistent comma structure
        first_line_commas = lines[0].count(',')
        if first_line_commas == 0:
            return False
            
        # Check if at least 80% of lines have similar comma count
        similar_lines = 0
        for line in lines[:min(5, len(lines))]:
            if abs(line.count(',') - first_line_commas) <= 1:
                similar_lines += 1
        
        return similar_lines / min(5, len(lines)) >= 0.8
    except:
        return False

def fetch_gmail_attachment(parent_message, attachment_key='attachment_0'):
    """
    Fetch Gmail attachment - simplified version that tries multiple approaches
    """
    try:
        logger.info(f"🔄 Attempting to fetch Gmail attachment for message {parent_message.get('id')}")
        
        # For now, return None - this will trigger the fallback content creation
        # In a production setup, this would integrate with Gmail API
        
        # TODO: Implement Gmail API integration when credentials are available
        # The function structure is ready for Gmail API implementation
        
        logger.warning(f"⚠️ Gmail API integration not configured - using fallback content")
        return None
        
    except Exception as e:
        logger.error(f"❌ Gmail attachment fetch error: {e}")
        return None

@app.route('/webhook/debug', methods=['POST'])
def debug_webhook():
    """Debug endpoint to see exactly what N8N is sending"""
    try:
        data = request.get_json()
        
        logger.info("📋 DEBUG WEBHOOK - Raw data received:")
        logger.info(f"Headers: {dict(request.headers)}")
        logger.info(f"Content-Type: {request.content_type}")
        logger.info(f"Data keys: {list(data.keys()) if data else 'No data'}")
        
        if data:
            for key, value in data.items():
                logger.info(f"  {key}: {type(value)} - {str(value)[:200]}...")
                
        return jsonify({
            'status': 'debug_received',
            'data_keys': list(data.keys()) if data else [],
            'has_binary': data.get('hasBinary', False) if data else False,
            'has_binary_data': data.get('hasBinaryData', False) if data else False,
            'binary_key': data.get('metadata', {}).get('binaryKey') if data else None,
            'content_preview': data.get('content', '')[:100] if data else None
        })
        
    except Exception as e:
        logger.error(f"❌ Debug webhook error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<document_id>', methods=['GET'])
def download_document(document_id):
    """Download original document file"""
    try:
        # Get document from database
        result = supabase.table('documents').select('*').eq('id', document_id).execute()
        
        if not result.data:
            return jsonify({'error': 'Document not found'}), 404
            
        document = result.data[0]
        
        # Check if binary data is available
        if not document.get('binary_data'):
            return jsonify({'error': 'Binary data not available for this document'}), 404
        
        try:
            # Decode base64 binary data
            import base64
            binary_content = base64.b64decode(document['binary_data'])
            
            # Determine content type
            filename = document.get('title', 'document')
            file_extension = filename.lower().split('.')[-1] if '.' in filename else 'txt'
            
            content_type_map = {
                'pdf': 'application/pdf',
                'csv': 'text/csv',
                'txt': 'text/plain',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'doc': 'application/msword',
                'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'xls': 'application/vnd.ms-excel'
            }
            
            content_type = content_type_map.get(file_extension, 'application/octet-stream')
            
            # Create response
            from flask import Response
            response = Response(
                binary_content,
                headers={
                    'Content-Type': content_type,
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'Content-Length': str(len(binary_content))
                }
            )
            
            logger.info(f"📥 Document downloaded: {filename} ({len(binary_content)} bytes)")
            return response
            
        except Exception as e:
            logger.error(f"❌ Error decoding binary data: {e}")
            return jsonify({'error': 'Error processing binary data'}), 500
            
    except Exception as e:
        logger.error(f"❌ Download error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/webhook/store-document', methods=['POST'])
def store_document_webhook():
    """Webhook endpoint for N8N document processing with actual binary file processing"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        logger.info(f"📥 Webhook document received from N8N")
        logger.info(f"🔍 Debug - Data keys: {list(data.keys())}")
        logger.info(f"🔍 Debug - Has binaryData: {bool(data.get('binaryData'))}")
        logger.info(f"🔍 Debug - Has binary object: {bool(data.get('binary'))}")
        logger.info(f"🔍 Debug - Has hasBinary flag: {data.get('hasBinary', False)}")
        
        # Extract document info
        document_id = data.get('documentId', '')
        filename = data.get('filename', 'unknown')
        content = data.get('content', '')
        source = data.get('source', 'n8n')
        metadata = data.get('metadata', {})
        parent_message = data.get('parentMessage', {})
        binary_data = data.get('binaryData')  # N8N sends binary data here
        
        # Process actual binary file content if available
        extracted_text = ""
        binary_content = None  # Initialize here
        binary_key = metadata.get('binaryKey', 'attachment_0')
        has_binary = data.get('hasBinary', metadata.get('hasBinaryData', False))
        
        if has_binary:
            try:
                logger.info(f"📎 Processing binary attachment: {filename}")
                
                # N8N can send binary data in multiple ways - check all possibilities
                
                # Method 1: Direct binaryData field
                if data.get('binaryData'):
                    import base64
                    binary_content = base64.b64decode(data['binaryData'])
                    logger.info("Found binary data in binaryData field")
                
                # Method 2: Binary object with key
                elif 'binary' in data and binary_key in data['binary']:
                    import base64
                    binary_info = data['binary'][binary_key]
                    if 'data' in binary_info:
                        binary_content = base64.b64decode(binary_info['data'])
                        logger.info(f"Found binary data in binary.{binary_key}.data")
                
                # Method 3: Check if content is base64 encoded binary
                elif content and len(content) > 100 and content.replace(' ', '').replace('\\n', '').replace('\\r', '').replace('+', '').replace('/', '').replace('=', '').isalnum():
                    try:
                        import base64
                        binary_content = base64.b64decode(content)
                        logger.info("Detected base64 content in content field")
                    except:
                        pass
                
                # Extract text if we found binary content
                if binary_content:
                    file_extension = filename.lower().split('.')[-1] if '.' in filename else 'txt'
                    
                    if file_extension == 'csv':
                        extracted_text = extract_text_from_csv(binary_content)
                    elif file_extension in ['pdf']:
                        extracted_text = extract_text_from_pdf(binary_content)
                    elif file_extension in ['doc', 'docx']:
                        extracted_text = extract_text_from_word(binary_content)
                    elif file_extension in ['xls', 'xlsx']:
                        extracted_text = extract_text_from_excel(binary_content)
                    elif file_extension in ['txt', 'text']:
                        extracted_text = binary_content.decode('utf-8', errors='ignore')
                    else:
                        # Try as text first
                        try:
                            extracted_text = binary_content.decode('utf-8', errors='ignore')
                        except:
                            extracted_text = f"Binary file: {filename} (content extraction not supported for {file_extension})"
                    
                    logger.info(f"✅ Extracted {len(extracted_text)} characters from {filename}")
                else:
                    logger.warning(f"⚠️ Binary data indicated but not found for {filename}")
                    extracted_text = f"Binary file: {filename} (binary data not accessible)"
                
            except Exception as e:
                logger.error(f"❌ Error processing binary data: {e}")
                extracted_text = f"Error extracting content from {filename}: {str(e)}"
        
        # Use extracted text as primary content, fall back to Gmail API if needed
        if extracted_text:
            content = extracted_text
        # Enhanced content processing - try to extract real data from N8N email processing
        if not extracted_text and data.get('hasBinary'):
            logger.info(f"🔍 Attempting enhanced content extraction from N8N data")
            
            # Check if N8N processed the email content and we can extract attachment data
            enhanced_content = extract_content_from_n8n_data(data, parent_message, metadata)
            if enhanced_content:
                content = enhanced_content['text']
                if enhanced_content.get('binary_data'):
                    binary_content = enhanced_content['binary_data']
                    extracted_text = enhanced_content['text']
                logger.info(f"✅ Enhanced extraction successful: {len(content)} chars")
            else:
                logger.info(f"⚠️ Enhanced extraction failed - using Gmail API fallback")
                # Try Gmail API as fallback
                gmail_attachment = fetch_gmail_attachment(parent_message, metadata.get('binaryKey', 'attachment_0'))
                
                if gmail_attachment:
                    binary_content = gmail_attachment['content']
                    extracted_text = gmail_attachment['text']
                    content = extracted_text
                    logger.info(f"✅ Gmail API fallback successful: {len(binary_content)} bytes")
                else:
                    logger.warning(f"⚠️ All extraction methods failed")
        
        # If still no content, create rich content from email metadata
        if not content and data.get('hasBinary'):
            # Create rich content from email metadata only if no actual content extracted
            subject = parent_message.get('subject', 'No subject')
            from_addr = parent_message.get('from', 'Unknown sender')
            file_type = data.get('mimeType', 'unknown')
            
            content = f"Document: {filename}\\nFrom: {from_addr}\\nSubject: {subject}\\nFile Type: {file_type}\\nNote: Binary content could not be extracted automatically."
        
        # Log the content being analyzed
        logger.info(f"📄 Content for analysis: {content[:200]}...")
        logger.info(f"📁 Filename: {filename}")
        logger.info(f"📧 Email subject: {parent_message.get('subject', 'N/A')}")
        
        # Process through OpenRouter RAG for department routing
        try:
            # Use the correct classifier variable name
            if multi_dept_classifier:
                rag_result = multi_dept_classifier.analyze_document(content, {'filename': filename})
                
                if rag_result and 'departments' in rag_result and rag_result['departments']:
                    assigned_department = rag_result['departments'][0]
                    priority = rag_result.get('priority', 'normal')
                    confidence = rag_result.get('confidence', 50)
                    logger.info(f"🎯 RAG Analysis: {filename} → {assigned_department} (confidence: {confidence}%)")
                else:
                    assigned_department = 'General'
                    priority = metadata.get('priority', 'normal')
                    confidence = 50
                    logger.info(f"🎯 Fallback routing: {filename} → {assigned_department}")
            else:
                logger.warning(f"⚠️ Multi-department classifier not available")
                assigned_department = 'General'
                priority = metadata.get('priority', 'normal')
                confidence = 50
                
        except Exception as e:
            logger.warning(f"⚠️ RAG analysis failed: {e}")
            # Fallback: simple keyword-based routing for finance documents
            content_lower = content.lower()
            filename_lower = filename.lower()
            
            if 'finance' in filename_lower or 'budget' in content_lower or 'revenue' in content_lower or 'expense' in content_lower:
                assigned_department = 'Finance'
                priority = 'normal'
                confidence = 75
                logger.info(f"🎯 Keyword fallback: {filename} → Finance (finance keywords detected)")
            elif 'maintenance' in content_lower or 'engineering' in content_lower or 'repair' in content_lower or 'inspection' in content_lower:
                assigned_department = 'Engineering'
                priority = 'normal'
                confidence = 75
                logger.info(f"🎯 Keyword fallback: {filename} → Engineering (maintenance keywords detected)")
            elif 'hr' in filename_lower or 'employee' in content_lower or 'performance' in content_lower:
                assigned_department = 'HR'
                priority = 'normal'
                confidence = 75
                logger.info(f"🎯 Keyword fallback: {filename} → HR (HR keywords detected)")
            else:
                assigned_department = 'General'
                priority = metadata.get('priority', 'normal')
                confidence = 50
                logger.info(f"🎯 Default fallback: {filename} → General")
        
        # Store actual binary data for download if available
        stored_binary = None
        file_size = 0
        
        if binary_content:
            try:
                # Store binary data in a way that can be retrieved for download
                # Option 1: Store as base64 in metadata (for smaller files)
                # Option 2: Store in file system (for larger files)
                
                file_size = len(binary_content)
                
                if file_size < 1024 * 1024:  # Less than 1MB - store in metadata
                    import base64
                    stored_binary = base64.b64encode(binary_content).decode('utf-8')
                    logger.info(f"💾 Stored {file_size} bytes as base64 in metadata")
                else:
                    # For larger files, could implement file system storage
                    logger.info(f"📁 Large file {file_size} bytes - would need file system storage")
                    
            except Exception as e:
                logger.error(f"❌ Error storing binary data: {e}")
        
        # Store in Supabase database
        doc_data = {
            'id': document_id,
            'title': filename,
            'content': content,
            'source': source,
            'language': 'english',
            'assigned_department': assigned_department,
            'priority': priority,
            'confidence': confidence,
            'status': 'processed',
            'processing_status': 'completed',
            'file_size': file_size,
            'content_extracted': bool(extracted_text),
            'binary_data': stored_binary,  # Store actual binary data
            'metadata': {
                **metadata,
                'webhook_source': True,
                'parent_message': parent_message,
                'content_hash': calculate_content_hash(content),
                'has_binary_data': bool(stored_binary),
                'original_file_size': file_size,
                'extraction_method': 'n8n_webhook_v2',
                'ai_analysis': {
                    'summary': f'Document processed from {source}',
                    'priority': priority,
                    'confidence': confidence,
                    'department': assigned_department,
                    'rag_analysis': True
                }
            }
        }
        
        # Insert into Supabase
        result = supabase.table('documents').insert(doc_data).execute()
        
        if result.data:
            logger.info(f"✅ Document stored successfully: {filename} → {assigned_department}")
            return jsonify({
                'success': True,
                'document_id': document_id,
                'assigned_department': assigned_department,
                'priority': priority,
                'confidence': confidence,
                'message': f'Document {filename} successfully processed and routed to {assigned_department}'
            }), 200
        else:
            logger.error(f"❌ Failed to store document in database")
            return jsonify({'error': 'Failed to store document in database'}), 500
            
    except Exception as e:
        logger.error(f"❌ Webhook processing error: {e}")
        return jsonify({'error': f'Webhook processing failed: {str(e)}'}), 500

@app.route('/api/process-pending', methods=['POST'])
def process_pending_documents():
    """Stage 2: Process pending documents with RAG"""
    try:
        # Get pending documents from database (using status = 'pending')
        result = supabase.table('documents')\
            .select('*')\
            .eq('status', 'pending')\
            .order('created_at', desc=False)\
            .limit(10)\
            .execute()
        
        pending_docs = result.data if result.data else []
        
        if not pending_docs:
            return jsonify({
                "status": "success",
                "message": "No pending documents to process",
                "processed_count": 0
            })
        
        processed_count = 0
        failed_count = 0
        
        for doc in pending_docs:
            try:
                doc_id = doc['id']
                content = doc.get('content', '')
                title = doc.get('title', '')
                metadata = doc.get('metadata', {})
                
                if not content:
                    logger.warning(f"⚠️ No content for document {doc_id}")
                    continue
                
                logger.info(f"🔄 Processing document {doc_id}")
                
                # Mark as processing in metadata
                updated_metadata = {**metadata, 'processing_status': 'processing'}
                supabase.table('documents').update({
                    'metadata': updated_metadata,
                    'status': 'processing',
                    'updated_at': datetime.now().isoformat()
                }).eq('id', doc_id).execute()
                
                # Process with RAG if available
                if rag_classifier:
                    rag_result = rag_classifier.analyze_with_rag(
                        content,
                        title,
                        metadata.get('subject', title)
                    )
                    
                    # Update document with RAG results including enhanced analysis
                    final_metadata = {
                        **updated_metadata,
                        'processing_status': 'completed',
                        'ai_analysis': rag_result,  # Store complete RAG analysis for frontend
                        'processed_at': datetime.now().isoformat()
                    }
                    
                    supabase.table('documents').update({
                        'assigned_department': rag_result.get('department', 'Administration'),
                        'confidence': rag_result.get('confidence_score', 70),
                        'priority': rag_result.get('priority', 'normal'),
                        'status': 'processed',
                        'metadata': final_metadata,
                        'updated_at': datetime.now().isoformat()
                    }).eq('id', doc_id).execute()
                    
                    processed_count += 1
                    logger.info(f"✅ Processed document {doc_id} → {rag_result.get('department')} (Enhanced Analysis: {bool(rag_result.get('document_summary'))})")
                else:
                    # Fallback to smart router
                    routing_result = smart_router.route_document(content, doc)
                    
                    final_metadata = {
                        **updated_metadata,
                        'processing_status': 'completed',
                        'routing_analysis': routing_result,
                        'processed_at': datetime.now().isoformat()
                    }
                    
                    supabase.table('documents').update({
                        'assigned_department': routing_result['assigned_department'],
                        'confidence': routing_result['confidence'],
                        'status': 'processed',
                        'metadata': final_metadata,
                        'updated_at': datetime.now().isoformat()
                    }).eq('id', doc_id).execute()
                    
                    processed_count += 1
                    logger.info(f"✅ Processed document {doc_id} → {routing_result['assigned_department']}")
                    
            except Exception as e:
                failed_count += 1
                logger.error(f"❌ Failed to process document {doc.get('id', 'unknown')}: {e}")
                
                # Mark as failed
                try:
                    error_metadata = {**doc.get('metadata', {}), 'processing_status': 'failed', 'error': str(e)}
                    supabase.table('documents').update({
                        'status': 'failed',
                        'metadata': error_metadata,
                        'updated_at': datetime.now().isoformat()
                    }).eq('id', doc.get('id')).execute()
                except:
                    pass
        
        return jsonify({
            "status": "success",
            "message": f"Processed {processed_count} documents",
            "processed_count": processed_count,
            "failed_count": failed_count,
            "total_pending": len(pending_docs)
        })
        
    except Exception as e:
        logger.error(f"❌ Error processing pending documents: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/processing-queue', methods=['GET'])
def get_processing_queue():
    """Get current processing queue status"""
    try:
        # Get pending documents
        pending_result = supabase.table('documents')\
            .select('id, title, source, status, created_at, metadata')\
            .eq('status', 'pending')\
            .order('created_at', desc=False)\
            .execute()
        
        # Get processing documents
        processing_result = supabase.table('documents')\
            .select('id, title, source, status, created_at, metadata')\
            .eq('status', 'processing')\
            .order('created_at', desc=False)\
            .execute()
        
        # Get recently processed documents
        processed_result = supabase.table('documents')\
            .select('id, title, source, status, assigned_department, confidence, created_at, updated_at, metadata')\
            .in_('status', ['processed', 'failed'])\
            .order('updated_at', desc=True)\
            .limit(10)\
            .execute()
        
        # Build queue information
        queue_items = []
        
        # Add pending items
        for doc in pending_result.data or []:
            metadata = doc.get('metadata', {})
            queue_items.append({
                'id': doc['id'],
                'title': doc['title'],
                'source': doc['source'],
                'status': 'pending',
                'processing_status': metadata.get('processing_status', 'pending'),
                'created_at': doc['created_at'],
                'queue_position': 'Ready for processing'
            })
        
        # Add processing items
        for doc in processing_result.data or []:
            metadata = doc.get('metadata', {})
            queue_items.append({
                'id': doc['id'],
                'title': doc['title'],
                'source': doc['source'],
                'status': 'processing',
                'processing_status': metadata.get('processing_status', 'processing'),
                'created_at': doc['created_at'],
                'queue_position': 'Currently processing'
            })
        
        # Calculate stats
        total_docs = supabase.table('documents').select('id', count='exact').execute()
        pending_count = len(pending_result.data or [])
        processing_count = len(processing_result.data or [])
        processed_count = len([d for d in (processed_result.data or []) if d['status'] == 'processed'])
        failed_count = len([d for d in (processed_result.data or []) if d['status'] == 'failed'])
        
        stats = {
            'total_documents': total_docs.count or 0,
            'pending_count': pending_count,
            'processing_count': processing_count,
            'processed_count': processed_count,
            'failed_count': failed_count
        }
        
        return jsonify({
            "status": "success",
            "queue": queue_items,
            "recent_processed": processed_result.data or [],
            "stats": stats,
            "queue_size": pending_count + processing_count
        })
        
    except Exception as e:
        logger.error(f"❌ Error fetching processing queue: {e}")
        return jsonify({"error": str(e)}), 500

# New authentication and enhanced functionality endpoints
@app.route('/api/chat/assistant', methods=['POST'])
def ai_assistant_chat():
    """AI Assistant chat endpoint using OpenRouter"""
    try:
        data = request.get_json()
        message = data.get('message', '').strip()
        department = data.get('department', 'General')
        context = data.get('context', 'ai_assistant')
        
        if not message:
            return jsonify({"error": "Message is required"}), 400
        
        # Create AI assistant prompt
        system_prompt = f"""You are an AI assistant for the {department} department in KMRL (Kochi Metro Rail Limited). 
Your role is to help with document management, task coordination, and departmental workflows.

Context: You're helping with {context} operations.
Department: {department}

You can help with:
- Document upload and processing (Excel, PDF, Word, images)
- Searching through department documents
- Content verification in documents
- Task management and coordination
- Department-specific guidance

Be helpful, professional, and specific to the {department} department context.
If asked about document routing, explain how documents are automatically classified to appropriate departments.
Keep responses concise but informative."""

        # Use OpenRouter for intelligent response
        headers = {
            "Authorization": f"Bearer {openrouter_api_key}",
            "HTTP-Referer": "http://localhost:3001",
            "X-Title": "DOC.X Intelligent AI Assistant",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "openai/gpt-3.5-turbo",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": message}
            ],
            "temperature": 0.7,
            "max_tokens": 500
        }
        
        response = requests.post(OPENROUTER_URL, headers=headers, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']
            
            return jsonify({
                "response": ai_response,
                "department": department,
                "context": context
            })
        else:
            logger.error(f"OpenRouter API error: {response.status_code}")
            return jsonify({
                "response": f"I'm here to help with {department} department tasks. How can I assist you today?",
                "department": department,
                "context": context
            })
        
    except Exception as e:
        logger.error(f"❌ Error in AI assistant chat: {e}")
        return jsonify({
            "response": "I'm experiencing some technical difficulties. Please try again.",
            "error": str(e)
        }), 500

@app.route('/api/search', methods=['GET'])
def search_documents():
    """Global document search across all documents"""
    try:
        query = request.args.get('q', '').strip()
        doc_type = request.args.get('type', '')
        department = request.args.get('department', '')
        
        if not query:
            return jsonify({"documents": [], "total": 0})
        
        # Build search filters
        filters = []
        if doc_type:
            filters.append(f"document_type.eq.{doc_type}")
        if department:
            filters.append(f"department.eq.{department}")
        
        # Search in title, content, and source
        search_query = supabase.table('documents').select('*')
        
        if filters:
            for filter_condition in filters:
                field, operator, value = filter_condition.split('.', 2)
                if operator == 'eq':
                    search_query = search_query.eq(field, value)
        
        # Text search across multiple fields
        search_query = search_query.or_(f"title.ilike.%{query}%,content.ilike.%{query}%,source.ilike.%{query}%")
        
        result = search_query.order('created_at', desc=True).limit(50).execute()
        
        # Add relevance scoring based on match location
        documents = []
        for doc in result.data or []:
            score = 0
            if query.lower() in doc.get('title', '').lower():
                score += 10
            if query.lower() in doc.get('content', '').lower():
                score += 5
            if query.lower() in doc.get('source', '').lower():
                score += 3
            
            doc['relevance_score'] = score
            documents.append(doc)
        
        # Sort by relevance
        documents.sort(key=lambda x: x['relevance_score'], reverse=True)
        
        return jsonify({
            "documents": documents,
            "total": len(documents),
            "query": query
        })
        
    except Exception as e:
        logger.error(f"❌ Error searching documents: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/verify-content', methods=['POST'])
def verify_content():
    """Verify if content exists in any document using RAG"""
    try:
        data = request.get_json()
        content = data.get('content', '').strip()
        
        if not content:
            return jsonify({"error": "Content is required"}), 400
        
        # Search for exact or similar content in documents
        result = supabase.table('documents').select('*').ilike('content', f'%{content}%').execute()
        
        matches = []
        for doc in result.data or []:
            doc_content = doc.get('content', '')
            # Find the context around the match
            content_lower = content.lower()
            doc_content_lower = doc_content.lower()
            
            if content_lower in doc_content_lower:
                start_idx = max(0, doc_content_lower.find(content_lower) - 100)
                end_idx = min(len(doc_content), start_idx + len(content) + 200)
                context = doc_content[start_idx:end_idx]
                
                matches.append({
                    'document_id': doc['id'],
                    'title': doc['title'],
                    'department': doc.get('department', 'Unknown'),
                    'context': context,
                    'match_found': True,
                    'source': doc.get('source', '')
                })
        
        return jsonify({
            "verified": len(matches) > 0,
            "matches": matches,
            "total_matches": len(matches),
            "search_content": content
        })
        
    except Exception as e:
        logger.error(f"❌ Error verifying content: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def ai_upload_document():
    """Enhanced upload with duplicate detection and privacy options"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        department = request.form.get('department', 'General')
        description = request.form.get('description', '')
        privacy_mode = request.form.get('privacy', 'global')  # 'global' or 'private'
        user_id = request.form.get('user_id', 'anonymous')
        force_upload = request.form.get('force_upload', 'false').lower() == 'true'
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        file.save(temp_path)
        
        try:
            # Process the document
            content = ""
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                content = extract_text_from_pdf(temp_path)
            elif file_ext in ['doc', 'docx']:
                content = extract_text_from_word(temp_path)
            elif file_ext in ['xls', 'xlsx']:
                content = extract_text_from_excel(temp_path)
            elif file_ext in ['txt']:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_ext in ['csv']:
                import pandas as pd
                df = pd.read_csv(temp_path)
                content = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\\n" + df.head().to_string()
            else:
                return jsonify({"error": f"Unsupported file type: {file_ext}"}), 400
            
            # Check for duplicates
            if not force_upload:
                duplicate_check = check_document_exists(filename, content, user_id)
                
                if duplicate_check['is_duplicate']:
                    existing_doc = duplicate_check['existing_content_docs'][0]
                    return jsonify({
                        "error": "Duplicate document detected",
                        "duplicate_info": {
                            "existing_title": existing_doc['title'],
                            "existing_id": existing_doc['id'],
                            "existing_department": existing_doc.get('assigned_department'),
                            "is_private": bool(existing_doc.get('private_to_user')),
                            "created_at": existing_doc.get('created_at')
                        },
                        "message": "A document with identical content already exists. Use force_upload=true to upload anyway.",
                        "force_upload_option": True
                    }), 409  # Conflict status code
                
                elif duplicate_check['exists_by_name']:
                    logger.warning(f"Document with same filename exists but different content: {filename}")
            
            # Use OpenRouter to intelligently route document to appropriate department
            if content.strip():
                try:
                    # Analyze document for department routing
                    if multi_dept_classifier:
                        routing_result = multi_dept_classifier.analyze_document(content, {'filename': filename})
                        
                        if routing_result and 'departments' in routing_result:
                            # Get the primary department (first one with highest confidence)
                            primary_dept = routing_result['departments'][0] if routing_result['departments'] else department
                            
                            # Use the AI-determined department instead of the uploaded department
                            routed_department = primary_dept
                            
                            logger.info(f"📋 Document '{filename}' auto-routed to {routed_department} department using OpenRouter RAG")
                        else:
                            routed_department = department
                            logger.info(f"📋 Document '{filename}' using fallback department: {department}")
                    else:
                        routed_department = department
                except Exception as routing_error:
                    logger.warning(f"⚠️ Department routing failed, using manual selection: {routing_error}")
                    routed_department = department
            else:
                routed_department = department
            
            # Calculate content hash for future duplicate detection
            content_hash = calculate_content_hash(content)
            
            # Store in database with privacy and duplicate detection
            doc_data = {
                'title': filename,
                'content': content,
                'source': f"AI Upload - {routed_department}",
                'document_type': file_ext,
                'assigned_department': routed_department,
                'status': 'processed',
                'private_to_user': user_id if privacy_mode == 'private' else None,
                'metadata': {
                    'uploaded_via': 'ai_interface',
                    'description': description,
                    'processing_status': 'completed',
                    'original_department_selection': department,
                    'ai_routed_department': routed_department,
                    'routing_method': 'OpenRouter RAG',
                    'privacy_mode': privacy_mode,
                    'file_size': len(content),
                    'duplicate_check_performed': True,
                    'force_upload': force_upload,
                    'content_hash': content_hash  # Store in metadata
                }
            }
            
            result = supabase.table('documents').insert(doc_data).execute()
            
            privacy_message = "privately (only visible to you)" if privacy_mode == 'private' else "globally (visible to all departments)"
            
            return jsonify({
                "success": True,
                "document_id": result.data[0]['id'],
                "message": f"Document '{filename}' uploaded {privacy_message} and intelligently routed to {routed_department} department using AI analysis",
                "filename": filename,
                "department": routed_department,
                "original_department": department,
                "routing_method": "OpenRouter RAG",
                "privacy_mode": privacy_mode,
                "content_hash": content_hash,
                "duplicate_check": "passed"
            })
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        logger.error(f"❌ Error uploading document: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/private-documents', methods=['GET', 'POST'])
def manage_private_documents():
    """Manage department-specific private documents"""
    try:
        if request.method == 'GET':
            department = request.args.get('department')
            if not department:
                return jsonify({"error": "Department is required"}), 400
            
            # Get private documents for department
            result = supabase.table('private_documents').select('*').eq('department', department).order('created_at', desc=True).execute()
            
            return jsonify({
                "documents": result.data or [],
                "total": len(result.data or []),
                "department": department
            })
        
        elif request.method == 'POST':
            if 'file' not in request.files:
                return jsonify({"error": "No file provided"}), 400
            
            file = request.files['file']
            department = request.form.get('department')
            tags = request.form.get('tags', '')
            is_confidential = request.form.get('is_confidential', 'false').lower() == 'true'
            
            if not department:
                return jsonify({"error": "Department is required"}), 400
            
            if file.filename == '':
                return jsonify({"error": "No file selected"}), 400
            
            filename = secure_filename(file.filename)
            
            # Create private documents table if not exists
            try:
                supabase.table('private_documents').select('id').limit(1).execute()
            except:
                # Table might not exist, create it
                pass
            
            # Store private document
            doc_data = {
                'title': filename,
                'department': department,
                'tags': tags.split(',') if tags else [],
                'is_confidential': is_confidential,
                'file_size': len(file.read()),
                'file_type': filename.split('.')[-1].lower(),
                'metadata': {
                    'original_filename': filename,
                    'upload_source': 'private_upload'
                }
            }
            
            file.seek(0)  # Reset file pointer
            
            result = supabase.table('private_documents').insert(doc_data).execute()
            
            return jsonify({
                "success": True,
                "document_id": result.data[0]['id'],
                "message": f"Private document '{filename}' added to {department}",
                "filename": filename
            })
    
    except Exception as e:
        logger.error(f"❌ Error managing private documents: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/department-summary', methods=['GET'])
def get_department_summary():
    """Get summary of all departments and their activities"""
    try:
        # Get document counts by department
        departments_result = supabase.table('documents').select('department').execute()
        
        department_stats = {}
        for doc in departments_result.data or []:
            dept = doc.get('department', 'Unknown')
            if dept not in department_stats:
                department_stats[dept] = {'total': 0, 'recent': 0}
            department_stats[dept]['total'] += 1
        
        # Get recent activity (last 7 days)
        from datetime import datetime, timedelta
        week_ago = (datetime.now() - timedelta(days=7)).isoformat()
        
        recent_result = supabase.table('documents').select('department').gte('created_at', week_ago).execute()
        
        for doc in recent_result.data or []:
            dept = doc.get('department', 'Unknown')
            if dept in department_stats:
                department_stats[dept]['recent'] += 1
        
        # Get pending tasks by department
        pending_result = supabase.table('documents').select('department').eq('status', 'pending').execute()
        
        for doc in pending_result.data or []:
            dept = doc.get('department', 'Unknown')
            if dept not in department_stats:
                department_stats[dept] = {'total': 0, 'recent': 0, 'pending': 0}
            if 'pending' not in department_stats[dept]:
                department_stats[dept]['pending'] = 0
            department_stats[dept]['pending'] += 1
        
        # Format response
        summaries = []
        for dept, stats in department_stats.items():
            summaries.append({
                'department': dept,
                'total_documents': stats['total'],
                'recent_activity': stats['recent'],
                'pending_tasks': stats.get('pending', 0),
                'status': 'active' if stats['recent'] > 0 else 'idle'
            })
        
        return jsonify({
            "department_summaries": summaries,
            "total_departments": len(summaries),
            "overall_stats": {
                "total_documents": sum(s['total_documents'] for s in summaries),
                "total_recent": sum(s['recent_activity'] for s in summaries),
                "total_pending": sum(s['pending_tasks'] for s in summaries)
            }
        })
        
    except Exception as e:
        logger.error(f"❌ Error getting department summary: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/tasks/<task_id>/complete', methods=['POST'])
def complete_task(task_id):
    """Mark a task as completed"""
    try:
        data = request.get_json()
        department = data.get('department')
        notes = data.get('notes', '')
        
        # Update document status
        result = supabase.table('documents').update({
            'status': 'completed',
            'metadata': {
                'completed_at': datetime.now().isoformat(),
                'completed_by': department,
                'completion_notes': notes
            }
        }).eq('id', task_id).execute()
        
        if result.data:
            return jsonify({
                "success": True,
                "message": "Task marked as completed",
                "task_id": task_id
            })
        else:
            return jsonify({"error": "Task not found"}), 404
        
    except Exception as e:
        logger.error(f"❌ Error completing task: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("=" * 60)
    print("🚀 DOC.X Intelligent Backend Starting...")
    print("=" * 60)
    print("📊 Features:")
    print("   • Multi-format document processing")
    print("   • Bilingual support (English/Malayalam)")
    print("   • Smart department routing")
    print("   • AI-powered analysis via OpenRouter")
    print("   • Supabase database integration")
    print("   • Two-stage processing: Storage → RAG")
    print("=" * 60)
    print("🌐 Endpoints:")
    print("   • Health Check: GET /health")
    print("   • Stage 1 Storage: POST /webhook/store-document")
    print("   • Document Processing: POST /webhook/document")
    print("   • Get Documents: GET /api/documents")
    print("   • Department Stats: GET /api/departments/stats")
    print("   • Chat Interface: POST /api/chat")
    print("   • RAG Stats: GET /api/rag/stats")
    print("   • RAG Retrain: POST /api/rag/retrain")
    print("   • Document Search: GET /api/search")
    print("   • Content Verification: POST /api/verify-content")
    print("   • AI Upload: POST /api/upload")
    print("   • Private Documents: GET/POST /api/private-documents")
    print("   • Department Summary: GET /api/department-summary")
    print("   • Task Completion: POST /api/tasks/<id>/complete")
    print("=" * 60)
    print("🎯 Ready for KMRL document processing!")
    print("=" * 60)
    
    app.run(host='0.0.0.0', port=5000, debug=False)