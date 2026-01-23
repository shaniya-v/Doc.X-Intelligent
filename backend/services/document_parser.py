"""
Document Parser Service
Extracts text from various document formats
"""
import logging
from typing import Optional
import io

# Document processing libraries
import PyPDF2
import pdfplumber
from PIL import Image
import pytesseract
from openpyxl import load_workbook
import docx
import pandas as pd

logger = logging.getLogger(__name__)

class DocumentParser:
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls', '.csv'],
        'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
        'text': ['.txt']
    }
    
    async def parse_document(self, filename: str, content: bytes) -> str:
        """
        Parse document based on file type
        Returns extracted text content
        """
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if f'.{file_extension}' in self.SUPPORTED_FORMATS['pdf']:
                return await self._parse_pdf(content)
            elif f'.{file_extension}' in self.SUPPORTED_FORMATS['docx']:
                return await self._parse_docx(content)
            elif f'.{file_extension}' in self.SUPPORTED_FORMATS['excel']:
                return await self._parse_excel(content, file_extension)
            elif f'.{file_extension}' in self.SUPPORTED_FORMATS['image']:
                return await self._parse_image(content)
            elif f'.{file_extension}' in self.SUPPORTED_FORMATS['text']:
                return content.decode('utf-8')
            else:
                logger.warning(f"Unsupported format: {file_extension}")
                return f"[Unsupported file format: {file_extension}]"
                
        except Exception as e:
            logger.error(f"❌ Parse error for {filename}: {e}", exc_info=True)
            return f"[Error parsing document: {str(e)}]"
    
    async def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            text_parts = []
            
            # Try pdfplumber first (better for tables)
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
                
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                logger.warning("No text extracted from PDF, might be image-based")
                return "[PDF contains no extractable text - might be scanned]"
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise
    
    async def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise
    
    async def _parse_excel(self, content: bytes, file_extension: str) -> str:
        """Extract text from Excel/CSV"""
        try:
            if file_extension == 'csv':
                df = pd.read_csv(io.BytesIO(content))
            else:
                df = pd.read_excel(io.BytesIO(content), sheet_name=None)
                
                # Combine all sheets
                if isinstance(df, dict):
                    all_data = []
                    for sheet_name, sheet_df in df.items():
                        all_data.append(f"Sheet: {sheet_name}")
                        all_data.append(sheet_df.to_string())
                    return "\n\n".join(all_data)
                
            return df.to_string()
            
        except Exception as e:
            logger.error(f"Excel parsing error: {e}")
            raise
    
    async def _parse_image(self, content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return "[No text found in image]"
            
            return text
            
        except Exception as e:
            logger.error(f"Image parsing error: {e}")
            raise
